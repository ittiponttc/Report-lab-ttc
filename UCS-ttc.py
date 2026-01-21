import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib import font_manager
import io
from datetime import datetime, date
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.chart import ScatterChart, Reference, Series
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64

# Page config
st.set_page_config(
    page_title="UCS Test Report Generator",
    page_icon="🧪",
    layout="wide"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2rem;
        font-weight: bold;
        color: #1e3a5f;
        text-align: center;
        padding: 1rem;
        background: linear-gradient(90deg, #e8f4f8 0%, #d1e8f0 100%);
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #2c5282;
        border-bottom: 2px solid #3182ce;
        padding-bottom: 0.5rem;
        margin-top: 1.5rem;
    }
    .info-box {
        background-color: #f0f9ff;
        border-left: 4px solid #3182ce;
        padding: 1rem;
        border-radius: 0 8px 8px 0;
        margin: 1rem 0;
    }
    .result-box {
        background-color: #f0fff4;
        border: 2px solid #48bb78;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
    }
    .stDownloadButton > button {
        width: 100%;
        background-color: #3182ce;
        color: white;
    }
</style>
""", unsafe_allow_html=True)

# Title
st.markdown('<div class="main-header">🧪 Unconfined Compression Test (ASTM D2166)<br>Report Generator</div>', unsafe_allow_html=True)
st.markdown('<p style="text-align: center; color: #666;">King Mongkut\'s University of Technology North Bangkok</p>', unsafe_allow_html=True)

# Initialize session state
if 'calculated' not in st.session_state:
    st.session_state.calculated = False
if 'results' not in st.session_state:
    st.session_state.results = None

# Sidebar for project information
st.sidebar.markdown("## 📋 Project Information")

with st.sidebar:
    project_name = st.text_input("Project Name / ชื่อโครงการ", "")
    specimen_from = st.text_input("Specimen From / ตัวอย่างจาก", "")
    location = st.text_input("Location / สถานที่", "")
    column_no = st.text_input("Column No. / หมายเลขเสาเข็ม", "")
    depth = st.text_input("Depth (m) / ความลึก", "7.50-8.50")
    sample_number = st.number_input("Sample Number / หมายเลขตัวอย่าง", min_value=1, value=1)
    
    st.markdown("---")
    st.markdown("### 📅 Dates")
    date_of_jetting = st.date_input("Date of Jetting / วันที่ฉีด", date.today())
    date_of_testing = st.date_input("Date of Testing / วันที่ทดสอบ", date.today())
    curing_time = st.number_input("Curing Time (days) / เวลาบ่ม", min_value=1, value=28)
    
    st.markdown("---")
    st.markdown("### 👤 Personnel")
    tested_by = st.text_input("Tested by / ผู้ทดสอบ", "Asst.Prof.Dr.Ittipon Meepon")
    controlled_by = st.text_input("Controlled Test by", "")
    notified_by = st.text_input("Notified by", "")

# Main content
col1, col2 = st.columns([1, 1])

with col1:
    st.markdown('<p class="sub-header">📊 Sample Properties</p>', unsafe_allow_html=True)
    
    subcol1, subcol2 = st.columns(2)
    with subcol1:
        diameter = st.number_input("Diameter, D (cm)", min_value=0.1, value=5.53, step=0.01, format="%.2f")
        height = st.number_input("Height, H (cm)", min_value=0.1, value=11.90, step=0.01, format="%.2f")
        weight = st.number_input("Weight of Sample, W (g)", min_value=0.1, value=621.07, step=0.01, format="%.2f")
    
    with subcol2:
        shearing_rate = st.text_input("Shearing Rate", "1% of sample height per minute")
        mixed_jet_mixing = st.number_input("Mixed Jet Mixing (kg/m³)", min_value=0.0, value=0.0, step=1.0)
    
    st.markdown('<p class="sub-header">💧 Water Content Determination</p>', unsafe_allow_html=True)
    
    wcol1, wcol2 = st.columns(2)
    with wcol1:
        container_no = st.text_input("Container No.", "9")
        weight_wet_soil_can = st.number_input("Weight of Wet Soil + Can (g)", min_value=0.0, value=406.14, step=0.01)
        weight_dry_soil_can = st.number_input("Weight of Dry Soil + Can (g)", min_value=0.0, value=244.31, step=0.01)
    with wcol2:
        weight_can = st.number_input("Weight of Can (g)", min_value=0.0, value=26.05, step=0.01)

with col2:
    st.markdown('<p class="sub-header">📁 Upload Test Data</p>', unsafe_allow_html=True)
    
    st.markdown("""
    <div class="info-box">
    <strong>📌 Required Format:</strong><br>
    Excel file with columns: <code>Load (kg)</code> and <code>Vertical displacement (mm)</code>
    </div>
    """, unsafe_allow_html=True)
    
    uploaded_file = st.file_uploader("Upload Excel File", type=['xlsx', 'xls'])
    
    proving_ring_capacity = st.number_input("Proving Ring Capacity (kN)", min_value=0.1, value=20.0, step=0.1)
    factor_k = st.number_input("Factor K (kg/division)", min_value=0.0001, value=1.8661, step=0.0001, format="%.4f")

# Calculate derived values
area = np.pi * (diameter / 2) ** 2  # cm²
volume = area * height  # cm³
wet_unit_weight = weight / volume if volume > 0 else 0  # g/cm³

# Water content calculation
weight_water = weight_wet_soil_can - weight_dry_soil_can
weight_dry_soil = weight_dry_soil_can - weight_can
water_content = (weight_water / weight_dry_soil * 100) if weight_dry_soil > 0 else 0
dry_unit_weight = wet_unit_weight / (1 + water_content/100) if water_content > 0 else wet_unit_weight

# Display calculated properties
st.markdown('<p class="sub-header">📐 Calculated Sample Properties</p>', unsafe_allow_html=True)
prop_col1, prop_col2, prop_col3, prop_col4 = st.columns(4)
with prop_col1:
    st.metric("Area, A (cm²)", f"{area:.2f}")
with prop_col2:
    st.metric("Volume, V (cm³)", f"{volume:.2f}")
with prop_col3:
    st.metric("Wet Unit Weight (g/cm³)", f"{wet_unit_weight:.2f}")
with prop_col4:
    st.metric("Dry Unit Weight (g/cm³)", f"{dry_unit_weight:.2f}")

wc_col1, wc_col2, wc_col3 = st.columns(3)
with wc_col1:
    st.metric("Weight of Water (g)", f"{weight_water:.2f}")
with wc_col2:
    st.metric("Weight of Dry Soil (g)", f"{weight_dry_soil:.2f}")
with wc_col3:
    st.metric("Water Content, w (%)", f"{water_content:.2f}")

# Process uploaded data
if uploaded_file is not None:
    try:
        df = pd.read_excel(uploaded_file)
        
        # Check for required columns
        required_cols = ['Load (kg)', 'Vertical displacement (mm)']
        if not all(col in df.columns for col in required_cols):
            st.error(f"❌ Required columns not found. Please ensure file has: {required_cols}")
        else:
            st.success("✅ Data loaded successfully!")
            
            # Calculate stress-strain data
            df['Deformation, R (mm)'] = df['Vertical displacement (mm)']
            df['Axial Load, P (kg)'] = df['Load (kg)']
            df['Axial Strain, ξ (R/10H)'] = df['Deformation, R (mm)'] / (10 * height)
            df['Axial Strain (%)'] = df['Axial Strain, ξ (R/10H)'] * 100
            df['Corrected Area, Ac (cm²)'] = area / (1 - df['Axial Strain, ξ (R/10H)'])
            df['Axial Stress, σ (ksc)'] = (df['Axial Load, P (kg)'] / df['Corrected Area, Ac (cm²)'])
            
            # Find maximum stress (UCS)
            max_stress_idx = df['Axial Stress, σ (ksc)'].idxmax()
            ucs = df.loc[max_stress_idx, 'Axial Stress, σ (ksc)']
            failure_strain = df.loc[max_stress_idx, 'Axial Strain (%)']
            undrained_shear_strength = ucs / 2
            
            # Calculate Modulus of Elasticity (E50)
            # E50 = stress at 50% of max stress / strain at 50% of max stress
            half_ucs = ucs / 2
            df_before_peak = df[df.index <= max_stress_idx]
            idx_e50 = (df_before_peak['Axial Stress, σ (ksc)'] - half_ucs).abs().idxmin()
            strain_at_half = df.loc[idx_e50, 'Axial Strain (%)']
            modulus_e50 = (half_ucs / (strain_at_half / 100)) if strain_at_half > 0 else 0
            
            # Store results
            results = {
                'ucs': ucs,
                'undrained_shear_strength': undrained_shear_strength,
                'failure_strain': failure_strain,
                'modulus_e50': modulus_e50,
                'df': df,
                'area': area,
                'volume': volume,
                'wet_unit_weight': wet_unit_weight,
                'dry_unit_weight': dry_unit_weight,
                'water_content': water_content,
                'weight_water': weight_water,
                'weight_dry_soil': weight_dry_soil
            }
            st.session_state.results = results
            st.session_state.calculated = True
            
            # Display results
            st.markdown('<p class="sub-header">🎯 Test Results</p>', unsafe_allow_html=True)
            
            res_col1, res_col2, res_col3, res_col4 = st.columns(4)
            with res_col1:
                st.metric("Unconfined Compressive Strength, qᵤ (ksc)", f"{ucs:.2f}")
            with res_col2:
                st.metric("Undrained Shear Strength, sᵤ (ksc)", f"{undrained_shear_strength:.2f}")
            with res_col3:
                st.metric("Failure Strain, εf (%)", f"{failure_strain:.2f}")
            with res_col4:
                st.metric("Modulus of Elasticity, E₅₀ (ksc)", f"{modulus_e50:.2f}")
            
            # Data table
            st.markdown('<p class="sub-header">📋 Test Data Table</p>', unsafe_allow_html=True)
            
            display_df = df[['Deformation, R (mm)', 'Axial Load, P (kg)', 'Axial Strain, ξ (R/10H)', 
                            'Corrected Area, Ac (cm²)', 'Axial Strain (%)', 'Axial Stress, σ (ksc)']].copy()
            display_df.columns = ['R (mm)', 'P (kg)', 'ξ (R/10H)', 'Ac (cm²)', 'ε (%)', 'σ (ksc)']
            st.dataframe(display_df.style.format({
                'R (mm)': '{:.2f}',
                'P (kg)': '{:.2f}',
                'ξ (R/10H)': '{:.2f}',
                'Ac (cm²)': '{:.2f}',
                'ε (%)': '{:.2f}',
                'σ (ksc)': '{:.2f}'
            }), use_container_width=True, height=300)
            
            # Plot
            st.markdown('<p class="sub-header">📈 Stress-Strain Curve</p>', unsafe_allow_html=True)
            
            fig, ax = plt.subplots(figsize=(10, 7))
            ax.plot(df['Axial Strain (%)'], df['Axial Stress, σ (ksc)'], 'b-o', markersize=5, linewidth=1.5)
            ax.scatter([failure_strain], [ucs], color='red', s=100, zorder=5, label=f'Peak: ({failure_strain:.2f}%, {ucs:.2f} ksc)')
            
            ax.set_xlabel('Axial Strain, ε, %', fontsize=12)
            ax.set_ylabel('Axial Stress, σ, ksc', fontsize=12)
            ax.set_title('Stress-Strain Curve', fontsize=14, fontweight='bold')
            ax.grid(True, linestyle='--', alpha=0.7)
            ax.legend(loc='lower right')
            ax.set_xlim(0, max(df['Axial Strain (%)']) * 1.1)
            ax.set_ylim(0, ucs * 1.2)
            
            st.pyplot(fig)
            
            # Save plot to buffer for export
            img_buffer = io.BytesIO()
            fig.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
            img_buffer.seek(0)
            
            # Download section
            st.markdown('<p class="sub-header">💾 Download Reports</p>', unsafe_allow_html=True)
            
            download_col1, download_col2 = st.columns(2)
            
            with download_col1:
                # Excel Export
                def create_excel_report():
                    wb = Workbook()
                    ws = wb.active
                    ws.title = "UCS Test Report"
                    
                    # Styles
                    header_font = Font(bold=True, size=12)
                    title_font = Font(bold=True, size=14)
                    border = Border(
                        left=Side(style='thin'),
                        right=Side(style='thin'),
                        top=Side(style='thin'),
                        bottom=Side(style='thin')
                    )
                    header_fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
                    
                    # Title
                    ws.merge_cells('A1:H1')
                    ws['A1'] = "KING MONGKUT'S UNIVERSITY OF TECHNOLOGY NORTH BANGKOK"
                    ws['A1'].font = title_font
                    ws['A1'].alignment = Alignment(horizontal='center')
                    
                    ws.merge_cells('A2:H2')
                    ws['A2'] = "Unconfined Compression Test (ASTM D2166)"
                    ws['A2'].font = header_font
                    ws['A2'].alignment = Alignment(horizontal='center')
                    
                    # Project Info
                    row = 4
                    info_data = [
                        ("Project Name:", project_name),
                        ("Specimen From:", specimen_from),
                        ("Location:", location),
                        ("Column No.:", column_no),
                        ("Depth:", depth),
                        ("Sample Number:", sample_number),
                        ("Date of Jetting:", str(date_of_jetting)),
                        ("Date of Testing:", str(date_of_testing)),
                        ("Curing Time (days):", curing_time),
                        ("Tested by:", tested_by),
                    ]
                    
                    for label, value in info_data:
                        ws[f'A{row}'] = label
                        ws[f'A{row}'].font = Font(bold=True)
                        ws[f'B{row}'] = value
                        row += 1
                    
                    # Sample Properties
                    row += 1
                    ws[f'A{row}'] = "Sample Properties"
                    ws[f'A{row}'].font = title_font
                    row += 1
                    
                    sample_data = [
                        ("Diameter, D (cm)", diameter),
                        ("Height, H (cm)", height),
                        ("Area, A (cm²)", area),
                        ("Volume, V (cm³)", volume),
                        ("Weight of Sample, W (g)", weight),
                        ("Wet Unit Weight (g/cm³)", wet_unit_weight),
                        ("Dry Unit Weight (g/cm³)", dry_unit_weight),
                        ("Water Content, w (%)", water_content),
                    ]
                    
                    for label, value in sample_data:
                        ws[f'A{row}'] = label
                        ws[f'B{row}'] = f"{value:.2f}" if isinstance(value, float) else value
                        row += 1
                    
                    # Results
                    row += 1
                    ws[f'A{row}'] = "Test Results"
                    ws[f'A{row}'].font = title_font
                    row += 1
                    
                    results_data = [
                        ("Unconfined Compressive Strength, qᵤ (ksc)", ucs),
                        ("Undrained Shear Strength, sᵤ (ksc)", undrained_shear_strength),
                        ("Failure Strain, εf (%)", failure_strain),
                        ("Modulus of Elasticity, E₅₀ (ksc)", modulus_e50),
                    ]
                    
                    for label, value in results_data:
                        ws[f'A{row}'] = label
                        ws[f'B{row}'] = f"{value:.2f}"
                        row += 1
                    
                    # Test Data Sheet
                    ws2 = wb.create_sheet("Test Data")
                    
                    headers = ['R (mm)', 'P (kg)', 'ξ', 'Ac (cm²)', 'ε (%)', 'σ (ksc)']
                    for col, header in enumerate(headers, 1):
                        cell = ws2.cell(row=1, column=col, value=header)
                        cell.font = header_font
                        cell.fill = header_fill
                        cell.border = border
                        cell.alignment = Alignment(horizontal='center')
                    
                    for r_idx, row_data in enumerate(df[['Deformation, R (mm)', 'Axial Load, P (kg)', 
                                                         'Axial Strain, ξ (R/10H)', 'Corrected Area, Ac (cm²)',
                                                         'Axial Strain (%)', 'Axial Stress, σ (ksc)']].values, 2):
                        for c_idx, value in enumerate(row_data, 1):
                            cell = ws2.cell(row=r_idx, column=c_idx, value=round(value, 2))
                            cell.border = border
                            cell.alignment = Alignment(horizontal='center')
                    
                    # Adjust column widths
                    from openpyxl.utils import get_column_letter # Import เพิ่มเพื่อใช้แปลงเลขคอลัมน์เป็นตัวอักษร
                    
                    for ws_sheet in [ws, ws2]:
                        for column_cells in ws_sheet.columns:
                            length = max(len(str(cell.value) if cell.value else "") for cell in column_cells)
                            
                            # แก้ไข: ใช้ get_column_letter(cell.column) แทน cell.column_letter
                            # เพราะ MergedCell จะไม่มี attribute column_letter แต่มี .column (int) เสมอ
                            col_idx = column_cells[0].column
                            col_letter = get_column_letter(col_idx)
                            
                            ws_sheet.column_dimensions[col_letter].width = max(length + 2, 12)
                    
                    # Save to buffer
                    excel_buffer = io.BytesIO()
                    wb.save(excel_buffer)
                    excel_buffer.seek(0)
                    return excel_buffer
                
                excel_data = create_excel_report()
                st.download_button(
                    label="📥 Download Excel Report",
                    data=excel_data,
                    file_name=f"UCS_Report_{project_name}_{date_of_testing}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            
            with download_col2:
                # Word Export
                def create_word_report():
                    doc = Document()
                    
                    # Title
                    title = doc.add_paragraph()
                    title_run = title.add_run("KING MONGKUT'S UNIVERSITY OF TECHNOLOGY NORTH BANGKOK")
                    title_run.bold = True
                    title_run.font.size = Pt(14)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    subtitle = doc.add_paragraph()
                    sub_run = subtitle.add_run("Department of Teacher Training in Civil Engineering")
                    sub_run.font.size = Pt(11)
                    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    doc.add_paragraph()
                    
                    test_title = doc.add_paragraph()
                    test_run = test_title.add_run("Unconfined Compression Test (ASTM D2166)")
                    test_run.bold = True
                    test_run.font.size = Pt(12)
                    test_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    doc.add_paragraph()
                    
                    # Project Info Table
                    info_table = doc.add_table(rows=6, cols=4)
                    info_table.style = 'Table Grid'
                    
                    info_data = [
                        ("Project Name:", project_name, "Tested by:", tested_by),
                        ("Specimen From:", specimen_from, "Date of Jetting:", str(date_of_jetting)),
                        ("Location:", location, "Date of Testing:", str(date_of_testing)),
                        ("Column No.:", column_no, "Curing Time:", f"{curing_time} days"),
                        ("Depth:", depth, "Sample Number:", str(sample_number)),
                        ("Shearing Rate:", shearing_rate, "", ""),
                    ]
                    
                    for i, (l1, v1, l2, v2) in enumerate(info_data):
                        row = info_table.rows[i]
                        row.cells[0].text = l1
                        row.cells[1].text = str(v1)
                        row.cells[2].text = l2
                        row.cells[3].text = str(v2)
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.size = Pt(10) if cell.paragraphs[0].runs else None
                    
                    doc.add_paragraph()
                    
                    # Sample Properties
                    props_heading = doc.add_paragraph()
                    props_run = props_heading.add_run("Soil-Jet Mixing Sample Properties")
                    props_run.bold = True
                    
                    props_table = doc.add_table(rows=8, cols=4)
                    props_table.style = 'Table Grid'
                    
                    props_data = [
                        ("Diameter, D", f"{diameter:.2f}", "cm", ""),
                        ("Height, H", f"{height:.2f}", "cm", ""),
                        ("Area, A", f"{area:.2f}", "cm²", ""),
                        ("Volume, V", f"{volume:.2f}", "cm³", ""),
                        ("Weight of Sample, W", f"{weight:.2f}", "g", ""),
                        ("Wet Unit Weight, γ", f"{wet_unit_weight:.2f}", "g/cm³", ""),
                        ("Dry Unit Weight, γd", f"{dry_unit_weight:.2f}", "g/cm³", ""),
                        ("Water Content, w", f"{water_content:.2f}", "%", ""),
                    ]
                    
                    for i, (param, value, unit, note) in enumerate(props_data):
                        row = props_table.rows[i]
                        row.cells[0].text = param
                        row.cells[1].text = value
                        row.cells[2].text = unit
                        row.cells[3].text = note
                    
                    doc.add_paragraph()
                    
                    # Results
                    results_heading = doc.add_paragraph()
                    results_run = results_heading.add_run("Test Results")
                    results_run.bold = True
                    
                    results_table = doc.add_table(rows=4, cols=3)
                    results_table.style = 'Table Grid'
                    
                    results_data = [
                        ("Unconfined Compressive Strength, (qᵤ)", f"{ucs:.2f}", "ksc"),
                        ("Undrained Shear Strength, (sᵤ)", f"{undrained_shear_strength:.2f}", "ksc"),
                        ("Failure Strain, (εf)", f"{failure_strain:.2f}", "%"),
                        ("Modulus of Elasticity, E₅₀", f"{modulus_e50:.2f}", "ksc"),
                    ]
                    
                    for i, (param, value, unit) in enumerate(results_data):
                        row = results_table.rows[i]
                        row.cells[0].text = param
                        row.cells[1].text = value
                        row.cells[2].text = unit
                    
                    doc.add_paragraph()
                    
                    # Add chart image
                    chart_heading = doc.add_paragraph()
                    chart_run = chart_heading.add_run("Stress-Strain Curve")
                    chart_run.bold = True
                    
                    doc.add_picture(img_buffer, width=Inches(5.5))
                    
                    doc.add_paragraph()
                    
                    # Data Table
                    data_heading = doc.add_paragraph()
                    data_run = data_heading.add_run("Test Data")
                    data_run.bold = True
                    
                    data_table = doc.add_table(rows=len(df)+1, cols=6)
                    data_table.style = 'Table Grid'
                    
                    headers = ['R (mm)', 'P (kg)', 'ξ', 'Ac (cm²)', 'ε (%)', 'σ (ksc)']
                    header_row = data_table.rows[0]
                    for i, header in enumerate(headers):
                        header_row.cells[i].text = header
                        header_row.cells[i].paragraphs[0].runs[0].bold = True
                    
                    for r_idx, row_data in enumerate(df[['Deformation, R (mm)', 'Axial Load, P (kg)', 
                                                         'Axial Strain, ξ (R/10H)', 'Corrected Area, Ac (cm²)',
                                                         'Axial Strain (%)', 'Axial Stress, σ (ksc)']].values):
                        row = data_table.rows[r_idx + 1]
                        for c_idx, value in enumerate(row_data):
                            row.cells[c_idx].text = f"{value:.2f}"
                    
                    doc.add_paragraph()
                    
                    # Remarks
                    remarks = doc.add_paragraph()
                    remarks_run = remarks.add_run("Remarks:")
                    remarks_run.bold = True
                    doc.add_paragraph("1. The testing results are good only for those specimens tested.")
                    doc.add_paragraph("2. Not valid unless signed and sealed.")
                    doc.add_paragraph(f"- Proving Ring Capacity: {proving_ring_capacity} kN")
                    doc.add_paragraph(f"- Factor K = {factor_k} kg/division")
                    
                    # Signature section
                    doc.add_paragraph()
                    sig_table = doc.add_table(rows=2, cols=3)
                    sig_table.rows[0].cells[0].text = "Test by:"
                    sig_table.rows[0].cells[1].text = "Controlled Test by:"
                    sig_table.rows[0].cells[2].text = "Notified by:"
                    sig_table.rows[1].cells[0].text = f"({tested_by})"
                    sig_table.rows[1].cells[1].text = f"({controlled_by})"
                    sig_table.rows[1].cells[2].text = f"({notified_by})"
                    
                    # Save to buffer
                    word_buffer = io.BytesIO()
                    doc.save(word_buffer)
                    word_buffer.seek(0)
                    return word_buffer
                
                word_data = create_word_report()
                st.download_button(
                    label="📥 Download Word Report",
                    data=word_data,
                    file_name=f"UCS_Report_{project_name}_{date_of_testing}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            plt.close()
            
    except Exception as e:
        st.error(f"❌ Error processing file: {str(e)}")
        st.exception(e)

else:
    st.info("👆 Please upload an Excel file with test data (Load and Vertical displacement) to generate the report.")

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #666; padding: 1rem;">
    <p>🏫 King Mongkut's University of Technology North Bangkok</p>
    <p>Department of Teacher Training in Civil Engineering</p>
    <p>© 2024 - UCS Test Report Generator</p>
</div>
""", unsafe_allow_html=True)
