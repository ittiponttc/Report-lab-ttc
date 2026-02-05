"""
Slope Stability & Settlement Analysis Application
วิเคราะห์เสถียรภาพลาดดินคันทางด้วยวิธี Bishop และ Swedish
พร้อมการวิเคราะห์การทรุดตัว

Developed for Civil Engineering Education
King Mongkut's University of Technology North Bangkok
"""

import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.patches import Arc, FancyArrowPatch
import json
from dataclasses import dataclass, asdict
from typing import List, Tuple, Optional
import math
from io import BytesIO
import warnings
warnings.filterwarnings('ignore')

# Configure matplotlib for Thai font support
plt.rcParams['font.family'] = 'DejaVu Sans'
plt.rcParams['axes.unicode_minus'] = False

# ============================================================
# Data Classes
# ============================================================

@dataclass
class SoilLayer:
    """ชั้นดิน"""
    name: str
    thickness: float  # m
    gamma: float  # kN/m³ (unit weight)
    gamma_sat: float  # kN/m³ (saturated unit weight)
    cohesion: float  # kPa
    phi: float  # degrees (friction angle)
    E: float  # kPa (Young's modulus)
    Cc: float  # Compression index
    Cr: float  # Recompression index
    e0: float  # Initial void ratio
    OCR: float  # Over-consolidation ratio
    Cv: float  # Coefficient of consolidation (m²/year)
    
    def __post_init__(self):
        """Ensure all numeric values are float"""
        self.thickness = float(self.thickness)
        self.gamma = float(self.gamma)
        self.gamma_sat = float(self.gamma_sat)
        self.cohesion = float(self.cohesion)
        self.phi = float(self.phi)
        self.E = float(self.E)
        self.Cc = float(self.Cc)
        self.Cr = float(self.Cr)
        self.e0 = float(self.e0)
        self.OCR = float(self.OCR)
        self.Cv = float(self.Cv)

@dataclass
class SlipCircle:
    """Slip circle parameters"""
    x_center: float
    y_center: float
    radius: float

@dataclass 
class AnalysisResults:
    """Analysis results"""
    method: str
    fs: float
    slices_data: List[dict]
    critical_circle: SlipCircle
    converged: bool = True
    iterations: int = 0

# ============================================================
# Slope Stability Analysis Functions
# ============================================================

def get_soil_at_point(x: float, y: float, slope_geometry: dict, soil_layers: List[SoilLayer], gwl: float) -> Tuple[Optional[SoilLayer], bool]:
    """
    หาชั้นดินที่จุด (x, y) และสถานะน้ำใต้ดิน
    """
    # Check if point is within slope
    slope_surface_y = get_slope_surface_y(x, slope_geometry)
    if y > slope_surface_y:
        return None, False
    
    # Find soil layer at this depth
    current_depth = slope_surface_y - y
    cumulative_thickness = 0
    
    for layer in soil_layers:
        cumulative_thickness += layer.thickness
        if current_depth <= cumulative_thickness:
            is_submerged = y < gwl
            return layer, is_submerged
    
    # Return last layer if below all defined layers
    if soil_layers:
        return soil_layers[-1], y < gwl
    return None, False

def get_slope_surface_y(x: float, slope_geometry: dict) -> float:
    """
    คำนวณความสูงผิวลาดดินที่ตำแหน่ง x
    """
    H = slope_geometry['height']
    slope_ratio = slope_geometry['slope_ratio']  # H:V ratio
    crest_width = slope_geometry['crest_width']
    toe_x = slope_geometry['toe_x']
    toe_elevation = slope_geometry.get('toe_elevation', 0.0)
    
    # Slope width
    slope_width = H * slope_ratio
    
    # Regions
    if x < toe_x:
        return toe_elevation  # Before toe (at toe elevation)
    elif x < toe_x + slope_width:
        # On slope
        return toe_elevation + (x - toe_x) / slope_ratio
    elif x < toe_x + slope_width + crest_width:
        return toe_elevation + H  # On crest
    else:
        return toe_elevation + H  # Beyond crest

def generate_slip_circle(slope_geometry: dict, x_offset: float = 0, y_offset: float = 0, r_factor: float = 1.0) -> SlipCircle:
    """
    สร้าง slip circle สำหรับการวิเคราะห์
    Center ต้องอยู่ด้านหน้า slope (เหนือ toe) เสมอ
    """
    H = slope_geometry['height']
    slope_ratio = slope_geometry['slope_ratio']
    toe_x = slope_geometry['toe_x']
    toe_elevation = slope_geometry.get('toe_elevation', 0.0)
    crest_width = slope_geometry.get('crest_width', 10.0)
    
    slope_width = H * slope_ratio
    crest_elevation = toe_elevation + H
    
    # Circle center should be in FRONT of slope (above/before toe)
    # X position: between toe and slightly before middle of slope
    x_center = toe_x + x_offset
    
    # Y position: above crest level
    y_center = crest_elevation + H * 0.3 + y_offset
    
    # Radius calculation - circle should pass through toe area and crest area
    # Distance from center to toe
    dist_to_toe = np.sqrt((x_center - toe_x)**2 + (y_center - toe_elevation)**2)
    
    # Distance from center to crest
    dist_to_crest = np.sqrt((x_center - (toe_x + slope_width))**2 + (y_center - crest_elevation)**2)
    
    # Use average distance as base radius
    base_radius = (dist_to_toe + dist_to_crest) / 2
    radius = base_radius * r_factor
    
    return SlipCircle(x_center, y_center, radius)

def slice_geometry(circle: SlipCircle, slope_geometry: dict, n_slices: int = 20) -> List[dict]:
    """
    แบ่ง slice สำหรับการวิเคราะห์
    """
    # Find intersection points with ground surface
    H = slope_geometry['height']
    toe_x = slope_geometry['toe_x']
    slope_ratio = slope_geometry['slope_ratio']
    toe_elevation = slope_geometry.get('toe_elevation', 0.0)
    slope_width = H * slope_ratio
    
    xc, yc, R = circle.x_center, circle.y_center, circle.radius
    
    # Find x range where circle intersects ground
    # Circle equation: (x-xc)² + (y-yc)² = R²
    # For y = toe_elevation: x = xc ± sqrt(R² - (yc-toe_elevation)²)
    
    y_diff = yc - toe_elevation
    if R**2 < y_diff**2:
        return []  # Circle doesn't reach ground level
    
    x_left = xc - np.sqrt(R**2 - y_diff**2)
    x_left = max(x_left, toe_x - 5)
    
    # Find right intersection (with slope surface) iteratively
    x_right = xc + np.sqrt(R**2 - y_diff**2)
    x_right = min(x_right, toe_x + slope_width + slope_geometry['crest_width'] + 5)
    
    # Create slices
    slice_width = (x_right - x_left) / n_slices
    slices = []
    
    for i in range(n_slices):
        x_mid = x_left + (i + 0.5) * slice_width
        x_l = x_left + i * slice_width
        x_r = x_left + (i + 1) * slice_width
        
        # Surface elevation at slice center
        y_surface = get_slope_surface_y(x_mid, slope_geometry)
        
        # Base of slice (on circle)
        y_base_sq = R**2 - (x_mid - xc)**2
        if y_base_sq < 0:
            continue
        y_base = yc - np.sqrt(y_base_sq)
        
        # Skip if base is above surface
        if y_base >= y_surface:
            continue
        
        # Slice height
        height = y_surface - y_base
        if height <= 0:
            continue
        
        # Base angle (alpha)
        alpha = np.arctan2(x_mid - xc, yc - y_base)
        
        slices.append({
            'index': i,
            'x_mid': x_mid,
            'x_left': x_l,
            'x_right': x_r,
            'y_surface': y_surface,
            'y_base': y_base,
            'height': height,
            'width': slice_width,
            'alpha': alpha,  # radians
            'alpha_deg': np.degrees(alpha)
        })
    
    return slices

def swedish_method(slices: List[dict], soil_layers: List[SoilLayer], 
                   slope_geometry: dict, gwl: float, circle: SlipCircle,
                   seismic_coef: float = 0.0) -> AnalysisResults:
    """
    วิเคราะห์ด้วยวิธี Swedish (Ordinary Method of Slices)
    FS = Σ(c'·l + (W·cos(α) - u·l)·tan(φ')) / Σ(W·sin(α) + kh·W·arm/R)
    
    For seismic (pseudo-static):
    - kh = horizontal seismic coefficient
    - Seismic force = kh * W (acting horizontally)
    - Additional driving moment = kh * W * (y_center - y_slice_center)
    """
    sum_resisting = 0
    sum_driving = 0
    slices_data = []
    
    R = circle.radius
    y_center = circle.y_center
    
    for s in slices:
        x_mid = s['x_mid']
        y_mid = (s['y_surface'] + s['y_base']) / 2
        
        # Get soil properties at slice center
        soil, is_submerged = get_soil_at_point(x_mid, y_mid, slope_geometry, soil_layers, gwl)
        if soil is None:
            continue
        
        # Use saturated unit weight if below GWL
        gamma = soil.gamma_sat if is_submerged else soil.gamma
        
        # Slice weight
        W = gamma * s['height'] * s['width']
        
        # Base length
        l = s['width'] / np.cos(s['alpha'])
        
        # Pore pressure at base
        if s['y_base'] < gwl:
            u = 9.81 * (gwl - s['y_base'])  # kPa
        else:
            u = 0
        
        # Normal and tangential forces
        alpha = s['alpha']
        N = W * np.cos(alpha) - u * l
        T = W * np.sin(alpha)
        
        # Seismic force contribution (pseudo-static)
        # Horizontal seismic force creates additional driving moment
        if seismic_coef > 0:
            # Moment arm for seismic force (horizontal distance from center)
            arm = y_center - y_mid  # Vertical distance from center to slice
            seismic_driving = seismic_coef * W * arm / R
            T += seismic_driving
        
        # Shear strength
        c = soil.cohesion
        phi_rad = np.radians(soil.phi)
        
        # Resisting force
        resisting = c * l + max(0, N) * np.tan(phi_rad)
        driving = T
        
        sum_resisting += resisting
        sum_driving += abs(driving) if driving > 0 else -driving
        
        slices_data.append({
            'index': s['index'],
            'x_mid': x_mid,
            'width': s['width'],
            'height': s['height'],
            'W': W,
            'alpha_deg': s['alpha_deg'],
            'l': l,
            'u': u,
            'N': N,
            'T': T,
            'c': c,
            'phi': soil.phi,
            'resisting': resisting,
            'driving': driving,
            'soil_name': soil.name
        })
    
    # Factor of Safety
    if abs(sum_driving) < 0.001:
        fs = 999
    else:
        fs = sum_resisting / abs(sum_driving)
    
    method_name = "Swedish (Ordinary Method)"
    if seismic_coef > 0:
        method_name += f" + Seismic (kh={seismic_coef})"
    
    return AnalysisResults(
        method=method_name,
        fs=fs,
        slices_data=slices_data,
        critical_circle=circle,
        converged=True,
        iterations=1
    )

def bishop_simplified(slices: List[dict], soil_layers: List[SoilLayer],
                      slope_geometry: dict, gwl: float, circle: SlipCircle,
                      seismic_coef: float = 0.0,
                      max_iter: int = 100, tol: float = 0.001) -> AnalysisResults:
    """
    วิเคราะห์ด้วยวิธี Bishop's Simplified Method
    FS = Σ[(c'·b + (W - u·b)·tan(φ')) / m_α] / Σ(W·sin(α) + kh·W·arm/R)
    where m_α = cos(α) + sin(α)·tan(φ')/FS
    
    For seismic (pseudo-static):
    - kh = horizontal seismic coefficient
    - Additional driving moment from horizontal seismic force
    """
    # Initial FS guess
    fs = 1.5
    
    slices_data = []
    R = circle.radius
    y_center = circle.y_center
    
    for iteration in range(max_iter):
        sum_numerator = 0
        sum_driving = 0
        temp_slices_data = []
        
        for s in slices:
            x_mid = s['x_mid']
            y_mid = (s['y_surface'] + s['y_base']) / 2
            
            # Get soil properties
            soil, is_submerged = get_soil_at_point(x_mid, y_mid, slope_geometry, soil_layers, gwl)
            if soil is None:
                continue
            
            gamma = soil.gamma_sat if is_submerged else soil.gamma
            
            # Slice weight
            W = gamma * s['height'] * s['width']
            b = s['width']
            
            # Pore pressure
            if s['y_base'] < gwl:
                u = 9.81 * (gwl - s['y_base'])
            else:
                u = 0
            
            alpha = s['alpha']
            c = soil.cohesion
            phi_rad = np.radians(soil.phi)
            
            # m_alpha factor
            m_alpha = np.cos(alpha) + np.sin(alpha) * np.tan(phi_rad) / fs
            
            # Prevent division by zero
            if abs(m_alpha) < 0.001:
                m_alpha = 0.001
            
            # Bishop equation terms
            numerator = (c * b + (W - u * b) * np.tan(phi_rad)) / m_alpha
            driving = W * np.sin(alpha)
            
            # Seismic force contribution (pseudo-static)
            if seismic_coef > 0:
                arm = y_center - y_mid
                seismic_driving = seismic_coef * W * arm / R
                driving += seismic_driving
            
            sum_numerator += numerator
            sum_driving += driving
            
            temp_slices_data.append({
                'index': s['index'],
                'x_mid': x_mid,
                'width': b,
                'height': s['height'],
                'W': W,
                'alpha_deg': s['alpha_deg'],
                'u': u,
                'm_alpha': m_alpha,
                'c': c,
                'phi': soil.phi,
                'numerator': numerator,
                'driving': driving,
                'soil_name': soil.name
            })
        
        # Calculate new FS
        if abs(sum_driving) < 0.001:
            fs_new = 999
        else:
            fs_new = sum_numerator / abs(sum_driving)
        
        # Check convergence
        if abs(fs_new - fs) < tol:
            slices_data = temp_slices_data
            method_name = "Bishop's Simplified"
            if seismic_coef > 0:
                method_name += f" + Seismic (kh={seismic_coef})"
            return AnalysisResults(
                method=method_name,
                fs=fs_new,
                slices_data=slices_data,
                critical_circle=circle,
                converged=True,
                iterations=iteration + 1
            )
        
        fs = fs_new
        slices_data = temp_slices_data
    
    method_name = "Bishop's Simplified"
    if seismic_coef > 0:
        method_name += f" + Seismic (kh={seismic_coef})"
    
    return AnalysisResults(
        method=method_name,
        fs=fs,
        slices_data=slices_data,
        critical_circle=circle,
        converged=False,
        iterations=max_iter
    )

def search_critical_circle(slope_geometry: dict, soil_layers: List[SoilLayer], 
                           gwl: float, method: str, n_circles: int = 50,
                           seismic_coef: float = 0.0) -> AnalysisResults:
    """
    ค้นหา critical slip circle ที่ให้ FS ต่ำสุด
    Circle center ต้องอยู่ด้านหน้า slope (เหนือ toe) เสมอตามหลักทฤษฎี
    
    Parameters:
    - seismic_coef: Horizontal seismic coefficient (kh) for pseudo-static analysis
    """
    H = slope_geometry['height']
    slope_ratio = slope_geometry['slope_ratio']
    toe_x = slope_geometry['toe_x']
    toe_elevation = slope_geometry.get('toe_elevation', 0.0)
    crest_width = slope_geometry.get('crest_width', 10.0)
    slope_width = H * slope_ratio
    crest_elevation = toe_elevation + H
    
    min_fs = float('inf')
    best_result = None
    
    # Grid search for circle center - CENTER MUST BE IN FRONT OF SLOPE
    # X range: from before toe to middle of slope
    x_range = np.linspace(toe_x - H*0.5, toe_x + slope_width * 0.3, int(np.sqrt(n_circles)))
    # Y range: from crest level to well above crest
    y_range = np.linspace(crest_elevation + H*0.2, crest_elevation + H*1.5, int(np.sqrt(n_circles)))
    r_factors = [0.8, 0.9, 1.0, 1.1, 1.2, 1.3, 1.4, 1.5]
    
    for x_c in x_range:
        for y_c in y_range:
            for r_f in r_factors:
                # Create circle directly with center position
                # Calculate radius to pass through slope
                dist_to_toe = np.sqrt((x_c - toe_x)**2 + (y_c - toe_elevation)**2)
                dist_to_crest = np.sqrt((x_c - (toe_x + slope_width))**2 + (y_c - crest_elevation)**2)
                base_radius = (dist_to_toe + dist_to_crest) / 2
                radius = base_radius * r_f
                
                circle = SlipCircle(x_c, y_c, radius)
                
                # Validate circle
                if radius < H * 0.5 or radius > H * 4:
                    continue
                
                # Circle bottom should reach below toe_elevation
                circle_bottom = y_c - radius
                if circle_bottom > toe_elevation - 0.5:
                    continue
                
                # Center must be above toe_elevation
                if y_c < toe_elevation + H * 0.5:
                    continue
                
                slices = slice_geometry(circle, slope_geometry, n_slices=15)
                
                if len(slices) < 5:
                    continue
                
                # Analyze with seismic if specified
                if method == "Swedish":
                    result = swedish_method(slices, soil_layers, slope_geometry, gwl, circle, seismic_coef)
                else:
                    result = bishop_simplified(slices, soil_layers, slope_geometry, gwl, circle, seismic_coef)
                
                if result.fs < min_fs and result.fs > 0.1:
                    min_fs = result.fs
                    best_result = result
    
    # Refine with more slices
    if best_result:
        circle = best_result.critical_circle
        slices = slice_geometry(circle, slope_geometry, n_slices=25)
        
        if method == "Swedish":
            best_result = swedish_method(slices, soil_layers, slope_geometry, gwl, circle, seismic_coef)
        else:
            best_result = bishop_simplified(slices, soil_layers, slope_geometry, gwl, circle, seismic_coef)
    
    return best_result

# ============================================================
# Settlement Analysis Functions
# ============================================================

def calculate_immediate_settlement(q: float, B: float, soil_layers: List[SoilLayer], 
                                   shape_factor: float = 1.0) -> float:
    """
    คำนวณการทรุดตัวทันที (Immediate Settlement)
    Si = q * B * (1 - ν²) * I / E
    """
    if not soil_layers:
        return 0
    
    # Use weighted average E
    total_thickness = sum(layer.thickness for layer in soil_layers)
    weighted_E = sum(layer.E * layer.thickness for layer in soil_layers) / total_thickness
    
    # Assume Poisson's ratio
    nu = 0.3
    
    # Influence factor (for flexible foundation)
    I = shape_factor * 0.88  # Corner of flexible rectangular
    
    Si = q * B * (1 - nu**2) * I / weighted_E
    
    return Si * 1000  # Convert to mm

def calculate_consolidation_settlement(q: float, soil_layers: List[SoilLayer],
                                       foundation_depth: float = 0) -> Tuple[float, List[dict]]:
    """
    คำนวณการทรุดตัวจากการอัดตัว (Consolidation Settlement)
    Sc = Σ [H * Cc * log10((σ'0 + Δσ) / σ'0) / (1 + e0)]
    """
    total_settlement = 0
    layer_settlements = []
    
    current_depth = foundation_depth
    sigma_v0 = 0  # Initial vertical effective stress
    
    for layer in soil_layers:
        # Mid-depth of layer
        mid_depth = current_depth + layer.thickness / 2
        
        # Effective stress at mid-layer
        gamma_eff = layer.gamma - 9.81  # Assume submerged
        sigma_v0 = mid_depth * gamma_eff
        
        if sigma_v0 <= 0:
            sigma_v0 = mid_depth * layer.gamma * 0.5
        
        # Stress increase (2:1 distribution approximation)
        delta_sigma = q / (1 + mid_depth / 5)**2
        
        # Preconsolidation pressure
        sigma_p = sigma_v0 * layer.OCR
        
        # Settlement calculation
        if sigma_v0 + delta_sigma <= sigma_p:
            # Recompression only
            Sc = (layer.thickness * layer.Cr / (1 + layer.e0)) * \
                 np.log10((sigma_v0 + delta_sigma) / sigma_v0)
        else:
            # Virgin compression
            if sigma_v0 < sigma_p:
                # Recompression part
                Sc_r = (layer.thickness * layer.Cr / (1 + layer.e0)) * \
                       np.log10(sigma_p / sigma_v0)
                # Virgin compression part
                Sc_v = (layer.thickness * layer.Cc / (1 + layer.e0)) * \
                       np.log10((sigma_v0 + delta_sigma) / sigma_p)
                Sc = Sc_r + Sc_v
            else:
                Sc = (layer.thickness * layer.Cc / (1 + layer.e0)) * \
                     np.log10((sigma_v0 + delta_sigma) / sigma_v0)
        
        total_settlement += Sc
        layer_settlements.append({
            'layer': layer.name,
            'thickness': layer.thickness,
            'sigma_v0': sigma_v0,
            'delta_sigma': delta_sigma,
            'settlement': Sc * 1000  # mm
        })
        
        current_depth += layer.thickness
    
    return total_settlement * 1000, layer_settlements  # Convert to mm

def calculate_time_rate_settlement(Sc: float, Cv: float, H_drainage: float, 
                                   times: List[float]) -> List[Tuple[float, float]]:
    """
    คำนวณอัตราการทรุดตัวตามเวลา (Time Rate of Settlement)
    U = 1 - Σ(2/M² * exp(-M²*Tv))
    where M = π(2m+1)/2, Tv = Cv*t/H²
    """
    results = []
    
    for t in times:
        if t <= 0:
            results.append((t, 0))
            continue
        
        # Time factor
        Tv = Cv * t / (H_drainage**2)
        
        # Degree of consolidation
        U = 0
        for m in range(50):
            M = np.pi * (2*m + 1) / 2
            U += (2 / M**2) * np.exp(-M**2 * Tv)
        U = 1 - U
        
        # Settlement at time t
        St = Sc * U
        results.append((t, St))
    
    return results

# ============================================================
# Visualization Functions
# ============================================================

def plot_slope_and_circle(slope_geometry: dict, soil_layers: List[SoilLayer],
                          gwl: float, result: AnalysisResults, 
                          show_slices: bool = True) -> plt.Figure:
    """
    วาดรูปลาดดินและ slip circle
    """
    fig, ax = plt.subplots(1, 1, figsize=(14, 10))
    
    H = slope_geometry['height']
    slope_ratio = slope_geometry['slope_ratio']
    toe_x = slope_geometry['toe_x']
    crest_width = slope_geometry['crest_width']
    toe_elevation = slope_geometry.get('toe_elevation', 0.0)
    slope_width = H * slope_ratio
    
    # Crest elevation
    crest_elevation = toe_elevation + H
    
    # Plot soil layers
    colors = plt.cm.YlOrBr(np.linspace(0.2, 0.8, len(soil_layers)))
    
    # Calculate total soil thickness and bottom elevation
    total_soil_thickness = sum(layer.thickness for layer in soil_layers)
    bottom_elevation = crest_elevation - total_soil_thickness
    
    # Draw ground/embankment fill
    slope_poly = plt.Polygon([
        (toe_x - 10, bottom_elevation - 2),  # Bottom left
        (toe_x - 10, toe_elevation),          # Left at toe level
        (toe_x, toe_elevation),               # Toe
        (toe_x + slope_width, crest_elevation),  # Top of slope
        (toe_x + slope_width + crest_width + 10, crest_elevation),  # Right at crest
        (toe_x + slope_width + crest_width + 10, bottom_elevation - 2)  # Bottom right
    ], facecolor='#8B7355', edgecolor='#5D4E37', linewidth=2)
    ax.add_patch(slope_poly)
    
    # Draw layer boundaries (from crest going down)
    cumulative = 0
    for i, layer in enumerate(soil_layers):
        y_top = crest_elevation - cumulative
        y_bottom = crest_elevation - cumulative - layer.thickness
        cumulative += layer.thickness
        
        ax.axhline(y=y_bottom, color='brown', linestyle='--', alpha=0.5, linewidth=1)
        ax.text(toe_x - 8, (y_top + y_bottom) / 2, 
                f'{layer.name}\nγ={layer.gamma:.1f}\nc={layer.cohesion:.1f}\nφ={layer.phi:.1f}°',
                fontsize=8, va='center', ha='left',
                bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.8))
    
    # Draw water table
    ax.axhline(y=gwl, color='blue', linestyle='-', linewidth=2, alpha=0.7, label=f'GWL = {gwl:.1f} m')
    ax.fill_between([toe_x - 10, toe_x + slope_width + crest_width + 10], 
                    bottom_elevation - 2, gwl, color='lightblue', alpha=0.3)
    
    # Draw slip circle
    if result and result.critical_circle:
        circle = result.critical_circle
        
        # Draw full circle (dashed)
        theta = np.linspace(0, 2*np.pi, 100)
        x_circle = circle.x_center + circle.radius * np.cos(theta)
        y_circle = circle.y_center + circle.radius * np.sin(theta)
        ax.plot(x_circle, y_circle, 'r--', linewidth=1, alpha=0.3)
        
        # Draw active portion (bottom half of circle)
        theta_range = np.linspace(-np.pi, 0, 100)
        x_active = circle.x_center + circle.radius * np.cos(theta_range)
        y_active = circle.y_center + circle.radius * np.sin(theta_range)
        
        # Filter to show only portion above bottom limit
        mask = y_active >= bottom_elevation - 1
        ax.plot(x_active[mask], y_active[mask], 'r-', linewidth=3, label='Slip Surface')
        
        # Mark center
        ax.plot(circle.x_center, circle.y_center, 'r+', markersize=15, markeredgewidth=3)
        ax.annotate(f'Center\n({circle.x_center:.1f}, {circle.y_center:.1f})',
                   (circle.x_center, circle.y_center),
                   textcoords="offset points", xytext=(10, 10),
                   fontsize=9, color='red')
    
    # Draw slices
    if show_slices and result and result.slices_data:
        for s in result.slices_data:
            x = s['x_mid']
            # Draw vertical line
            ax.axvline(x=x, ymin=0, ymax=0.9, color='green', 
                      linestyle=':', alpha=0.5, linewidth=0.5)
    
    # Add FS result
    if result:
        status_color = 'green' if result.fs >= 1.5 else 'orange' if result.fs >= 1.0 else 'red'
        status_text = 'SAFE' if result.fs >= 1.5 else 'MARGINAL' if result.fs >= 1.0 else 'UNSAFE'
        
        ax.text(0.98, 0.98, 
                f'{result.method}\nFS = {result.fs:.3f}\n{status_text}',
                transform=ax.transAxes, fontsize=14, fontweight='bold',
                verticalalignment='top', horizontalalignment='right',
                bbox=dict(boxstyle='round', facecolor=status_color, alpha=0.3))
        
        if not result.converged:
            ax.text(0.98, 0.78, 'Warning: Not Converged',
                   transform=ax.transAxes, fontsize=10, color='red',
                   verticalalignment='top', horizontalalignment='right')
    
    # Labels
    ax.set_xlabel('Distance (m)', fontsize=12)
    ax.set_ylabel('Elevation (m)', fontsize=12)
    ax.set_title('Slope Stability Analysis - Slip Circle Method', fontsize=14, fontweight='bold')
    ax.legend(loc='upper left')
    ax.grid(True, alpha=0.3)
    ax.set_aspect('equal')
    
    # Set limits - adjust for toe_elevation
    ax.set_xlim(toe_x - 12, toe_x + slope_width + crest_width + 12)
    ax.set_ylim(min(bottom_elevation - 3, toe_elevation - 5), crest_elevation + H * 0.5)
    
    plt.tight_layout()
    return fig

def plot_settlement_time(settlement_data: List[Tuple[float, float]], 
                         total_settlement: float) -> plt.Figure:
    """
    วาดกราฟการทรุดตัวตามเวลา
    """
    fig, ax = plt.subplots(1, 1, figsize=(10, 6))
    
    times = [d[0] for d in settlement_data]
    settlements = [d[1] for d in settlement_data]
    
    ax.plot(times, settlements, 'b-', linewidth=2, marker='o', markersize=4)
    ax.axhline(y=total_settlement, color='r', linestyle='--', 
               label=f'Ultimate Settlement = {total_settlement:.1f} mm')
    
    # Mark 50% and 90% consolidation
    for pct in [0.5, 0.9]:
        target = total_settlement * pct
        for i, s in enumerate(settlements):
            if s >= target:
                ax.axhline(y=target, color='gray', linestyle=':', alpha=0.5)
                ax.annotate(f'{int(pct*100)}% @ t={times[i]:.1f} yr',
                           (times[i], target), textcoords="offset points",
                           xytext=(10, 5), fontsize=9)
                break
    
    ax.set_xlabel('Time (years)', fontsize=12)
    ax.set_ylabel('Settlement (mm)', fontsize=12)
    ax.set_title('Consolidation Settlement vs Time', fontsize=14, fontweight='bold')
    ax.legend()
    ax.grid(True, alpha=0.3)
    ax.invert_yaxis()  # Settlement increases downward
    
    plt.tight_layout()
    return fig

def plot_stress_distribution(soil_layers: List[SoilLayer], q: float) -> plt.Figure:
    """
    วาดการกระจายของหน่วยแรงตามความลึก
    """
    fig, axes = plt.subplots(1, 3, figsize=(14, 8))
    
    depths = []
    sigma_v = []
    sigma_v0 = []
    delta_sigma = []
    
    current_depth = 0
    current_stress = 0
    
    for layer in soil_layers:
        n_points = 10
        for i in range(n_points):
            d = current_depth + layer.thickness * i / n_points
            depths.append(d)
            
            # Initial effective stress
            stress = current_stress + (layer.gamma - 9.81) * layer.thickness * i / n_points
            sigma_v0.append(max(stress, 0))
            
            # Stress increase (2:1 approximation)
            ds = q / (1 + d / 5)**2
            delta_sigma.append(ds)
            
            # Total effective stress
            sigma_v.append(max(stress, 0) + ds)
        
        current_depth += layer.thickness
        current_stress += (layer.gamma - 9.81) * layer.thickness
    
    # Plot initial stress
    axes[0].plot(sigma_v0, depths, 'b-', linewidth=2)
    axes[0].fill_betweenx(depths, 0, sigma_v0, alpha=0.3)
    axes[0].set_xlabel('σ\'₀ (kPa)', fontsize=11)
    axes[0].set_ylabel('Depth (m)', fontsize=11)
    axes[0].set_title('Initial Effective Stress', fontsize=12, fontweight='bold')
    axes[0].invert_yaxis()
    axes[0].grid(True, alpha=0.3)
    
    # Plot stress increase
    axes[1].plot(delta_sigma, depths, 'r-', linewidth=2)
    axes[1].fill_betweenx(depths, 0, delta_sigma, alpha=0.3, color='red')
    axes[1].set_xlabel('Δσ (kPa)', fontsize=11)
    axes[1].set_ylabel('Depth (m)', fontsize=11)
    axes[1].set_title('Stress Increase', fontsize=12, fontweight='bold')
    axes[1].invert_yaxis()
    axes[1].grid(True, alpha=0.3)
    
    # Plot final stress
    axes[2].plot(sigma_v0, depths, 'b--', linewidth=1.5, label='Initial')
    axes[2].plot(sigma_v, depths, 'g-', linewidth=2, label='Final')
    axes[2].fill_betweenx(depths, sigma_v0, sigma_v, alpha=0.3, color='green')
    axes[2].set_xlabel('σ\'v (kPa)', fontsize=11)
    axes[2].set_ylabel('Depth (m)', fontsize=11)
    axes[2].set_title('Effective Stress Profile', fontsize=12, fontweight='bold')
    axes[2].invert_yaxis()
    axes[2].legend()
    axes[2].grid(True, alpha=0.3)
    
    # Mark layers
    cum_depth = 0
    for layer in soil_layers:
        cum_depth += layer.thickness
        for ax in axes:
            ax.axhline(y=cum_depth, color='brown', linestyle=':', alpha=0.5)
    
    plt.tight_layout()
    return fig

# ============================================================
# JSON Save/Load Functions
# ============================================================

def save_to_json(slope_geometry: dict, soil_layers: List[SoilLayer], 
                 gwl: float, analysis_params: dict) -> str:
    """บันทึกข้อมูลเป็น JSON"""
    data = {
        'slope_geometry': slope_geometry,
        'soil_layers': [asdict(layer) for layer in soil_layers],
        'gwl': gwl,
        'analysis_params': analysis_params
    }
    return json.dumps(data, indent=2, ensure_ascii=False)

def load_from_json(json_str: str) -> Tuple[dict, List[SoilLayer], float, dict]:
    """โหลดข้อมูลจาก JSON"""
    data = json.loads(json_str)
    
    # Convert slope geometry values to float
    slope_geometry = {
        'height': float(data['slope_geometry']['height']),
        'slope_ratio': float(data['slope_geometry']['slope_ratio']),
        'crest_width': float(data['slope_geometry']['crest_width']),
        'toe_x': float(data['slope_geometry']['toe_x']),
        'toe_elevation': float(data['slope_geometry'].get('toe_elevation', 0.0))
    }
    
    # Convert soil layer values to float
    soil_layers = []
    for layer in data['soil_layers']:
        soil_layers.append(SoilLayer(
            name=layer['name'],
            thickness=float(layer['thickness']),
            gamma=float(layer['gamma']),
            gamma_sat=float(layer['gamma_sat']),
            cohesion=float(layer['cohesion']),
            phi=float(layer['phi']),
            E=float(layer['E']),
            Cc=float(layer['Cc']),
            Cr=float(layer['Cr']),
            e0=float(layer['e0']),
            OCR=float(layer['OCR']),
            Cv=float(layer['Cv'])
        ))
    
    gwl = float(data['gwl'])
    analysis_params = data.get('analysis_params', {})
    
    return slope_geometry, soil_layers, gwl, analysis_params

# ============================================================
# Word Report Export Function
# ============================================================

def generate_word_report(slope_geometry: dict, soil_layers: List[SoilLayer],
                        gwl: float, result: AnalysisResults, 
                        seismic_coef: float = 0.0,
                        fig_slope: plt.Figure = None) -> bytes:
    """
    สร้างรายงาน Word สำหรับการวิเคราะห์เสถียรภาพลาดดิน
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(14)
    
    # Title
    title = doc.add_heading('รายงานการวิเคราะห์เสถียรภาพลาดดิน', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Slope Stability Analysis Report', style='Heading 2').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Project Information
    doc.add_heading('1. ข้อมูลโครงการ (Project Information)', level=1)
    
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('Analysis Method', result.method if result else 'N/A'),
        ('Embankment Height (m)', f"{slope_geometry['height']:.2f}"),
        ('Slope Ratio (H:V)', f"{slope_geometry['slope_ratio']:.2f}:1"),
        ('Toe Elevation (m)', f"{slope_geometry.get('toe_elevation', 0):.2f}"),
        ('Ground Water Level (m)', f"{gwl:.2f}"),
        ('Seismic Coefficient (kh)', f"{seismic_coef:.3f}" if seismic_coef > 0 else "Not Applied"),
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # Soil Properties
    doc.add_heading('2. คุณสมบัติดิน (Soil Properties)', level=1)
    
    soil_table = doc.add_table(rows=len(soil_layers)+1, cols=7)
    soil_table.style = 'Table Grid'
    
    # Header
    headers = ['Layer', 'Thickness (m)', 'γ (kN/m³)', 'γsat (kN/m³)', 'c\' (kPa)', 'φ\' (°)', 'E (kPa)']
    for j, header in enumerate(headers):
        soil_table.rows[0].cells[j].text = header
    
    # Data
    for i, layer in enumerate(soil_layers):
        soil_table.rows[i+1].cells[0].text = layer.name
        soil_table.rows[i+1].cells[1].text = f"{layer.thickness:.2f}"
        soil_table.rows[i+1].cells[2].text = f"{layer.gamma:.1f}"
        soil_table.rows[i+1].cells[3].text = f"{layer.gamma_sat:.1f}"
        soil_table.rows[i+1].cells[4].text = f"{layer.cohesion:.1f}"
        soil_table.rows[i+1].cells[5].text = f"{layer.phi:.1f}"
        soil_table.rows[i+1].cells[6].text = f"{layer.E:.0f}"
    
    doc.add_paragraph()
    
    # Analysis Results
    doc.add_heading('3. ผลการวิเคราะห์ (Analysis Results)', level=1)
    
    if result:
        # Factor of Safety
        fs = result.fs
        status = 'SAFE (ปลอดภัย)' if fs >= 1.5 else 'MARGINAL (พอใช้)' if fs >= 1.0 else 'UNSAFE (ไม่ปลอดภัย)'
        
        result_para = doc.add_paragraph()
        result_para.add_run(f'Factor of Safety (FS) = {fs:.3f}').bold = True
        doc.add_paragraph(f'Status: {status}')
        
        if not result.converged:
            doc.add_paragraph('Warning: Analysis did not converge!').italic = True
        
        doc.add_paragraph(f'Number of iterations: {result.iterations}')
        
        # Critical Circle
        doc.add_paragraph()
        doc.add_heading('Critical Slip Circle Parameters:', level=2)
        circle = result.critical_circle
        doc.add_paragraph(f'• Center X: {circle.x_center:.2f} m')
        doc.add_paragraph(f'• Center Y: {circle.y_center:.2f} m')
        doc.add_paragraph(f'• Radius: {circle.radius:.2f} m')
        
        # Slice Data Table
        if result.slices_data:
            doc.add_paragraph()
            doc.add_heading('4. รายละเอียด Slice (Slice Details)', level=1)
            
            slice_table = doc.add_table(rows=min(len(result.slices_data)+1, 26), cols=6)
            slice_table.style = 'Table Grid'
            
            slice_headers = ['Slice', 'Width (m)', 'Height (m)', 'W (kN)', 'α (°)', 'Soil']
            for j, header in enumerate(slice_headers):
                slice_table.rows[0].cells[j].text = header
            
            for i, s in enumerate(result.slices_data[:25]):  # Limit to 25 slices
                slice_table.rows[i+1].cells[0].text = str(s['index'] + 1)
                slice_table.rows[i+1].cells[1].text = f"{s['width']:.2f}"
                slice_table.rows[i+1].cells[2].text = f"{s['height']:.2f}"
                slice_table.rows[i+1].cells[3].text = f"{s['W']:.1f}"
                slice_table.rows[i+1].cells[4].text = f"{s['alpha_deg']:.1f}"
                slice_table.rows[i+1].cells[5].text = s['soil_name']
    
    # Add figure if provided
    if fig_slope:
        doc.add_paragraph()
        doc.add_heading('5. แผนภาพการวิเคราะห์ (Analysis Diagram)', level=1)
        
        # Save figure to bytes
        img_buffer = io.BytesIO()
        fig_slope.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
        img_buffer.seek(0)
        
        doc.add_picture(img_buffer, width=Inches(6))
    
    # Conclusion
    doc.add_paragraph()
    doc.add_heading('6. สรุปผล (Conclusion)', level=1)
    
    if result:
        if result.fs >= 1.5:
            conclusion = f"""จากการวิเคราะห์เสถียรภาพลาดดินด้วยวิธี {result.method} 
            พบว่าค่า Factor of Safety (FS) = {result.fs:.3f} ซึ่งมากกว่า 1.5 
            ดังนั้นลาดดินมีความปลอดภัยเพียงพอตามมาตรฐานการออกแบบ"""
        elif result.fs >= 1.0:
            conclusion = f"""จากการวิเคราะห์เสถียรภาพลาดดินด้วยวิธี {result.method}
            พบว่าค่า Factor of Safety (FS) = {result.fs:.3f} ซึ่งอยู่ระหว่าง 1.0-1.5
            ลาดดินมีความเสถียรพอใช้ แต่ควรพิจารณาปรับปรุงหรือเพิ่มมาตรการเสริมความมั่นคง"""
        else:
            conclusion = f"""จากการวิเคราะห์เสถียรภาพลาดดินด้วยวิธี {result.method}
            พบว่าค่า Factor of Safety (FS) = {result.fs:.3f} ซึ่งน้อยกว่า 1.0
            ลาดดินไม่มีความปลอดภัย จำเป็นต้องปรับปรุงการออกแบบหรือเพิ่มมาตรการเสริมความมั่นคง"""
        
        doc.add_paragraph(conclusion)
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    doc.add_paragraph('Generated by Slope Stability Analysis Application')
    doc.add_paragraph('Department of Civil Engineering, KMUTNB')
    
    # Save to bytes
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer.getvalue()

# ============================================================
# Streamlit Application
# ============================================================

def main():
    st.set_page_config(
        page_title="Slope Stability & Settlement Analysis",
        page_icon="🏔️",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS
    st.markdown("""
    <style>
    .main-header {
        font-size: 2.2rem;
        font-weight: bold;
        color: #1E3A5F;
        text-align: center;
        padding: 1rem;
        background: linear-gradient(135deg, #E8F4F8 0%, #B8D4E3 100%);
        border-radius: 10px;
        margin-bottom: 2rem;
        border-left: 5px solid #1E3A5F;
    }
    .sub-header {
        font-size: 1.3rem;
        color: #2C5F7C;
        font-weight: 600;
        margin-top: 1.5rem;
        margin-bottom: 0.5rem;
        border-bottom: 2px solid #B8D4E3;
        padding-bottom: 0.3rem;
    }
    .info-box {
        background: linear-gradient(135deg, #F0F7FF 0%, #E8F4F8 100%);
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #1E3A5F;
        margin: 1rem 0;
    }
    .result-safe {
        background: linear-gradient(135deg, #D4EDDA 0%, #C3E6CB 100%);
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 5px solid #28A745;
        text-align: center;
    }
    .result-warning {
        background: linear-gradient(135deg, #FFF3CD 0%, #FFE69C 100%);
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 5px solid #FFC107;
        text-align: center;
    }
    .result-danger {
        background: linear-gradient(135deg, #F8D7DA 0%, #F5C6CB 100%);
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 5px solid #DC3545;
        text-align: center;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    .stTabs [data-baseweb="tab"] {
        background-color: #E8F4F8;
        border-radius: 8px 8px 0 0;
        padding: 10px 20px;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Header
    st.markdown('<div class="main-header">🏔️ Slope Stability & Settlement Analysis<br/><span style="font-size:1rem;">วิเคราะห์เสถียรภาพลาดดินคันทางและการทรุดตัว</span></div>', unsafe_allow_html=True)
    
    # Initialize session state
    if 'soil_layers' not in st.session_state:
        st.session_state.soil_layers = [
            SoilLayer("Layer 1 - Clay", 3.0, 18.0, 19.5, 25.0, 10.0, 15000, 0.35, 0.08, 0.8, 1.5, 1.0),
            SoilLayer("Layer 2 - Silty Clay", 4.0, 17.5, 19.0, 15.0, 18.0, 20000, 0.25, 0.06, 0.7, 2.0, 2.5),
        ]
    if 'analysis_result' not in st.session_state:
        st.session_state.analysis_result = None
    
    # Sidebar - Data Management
    with st.sidebar:
        st.markdown("### 📁 Data Management")
        
        # JSON Upload
        uploaded_file = st.file_uploader("Upload JSON Configuration", type=['json'])
        if uploaded_file:
            try:
                json_content = uploaded_file.read().decode('utf-8')
                slope_geo, layers, gwl_val, params = load_from_json(json_content)
                st.session_state.soil_layers = layers
                st.session_state.loaded_slope = slope_geo
                st.session_state.loaded_gwl = gwl_val
                st.success("✅ Configuration loaded successfully!")
            except Exception as e:
                st.error(f"Error loading file: {str(e)}")
        
        st.markdown("---")
        st.markdown("### 📚 Analysis Methods")
        st.markdown("""
        **Swedish Method (OMS)**
        - Simple, conservative
        - No iteration required
        - Suitable for preliminary analysis
        
        **Bishop's Simplified**
        - More accurate
        - Iterative solution
        - Industry standard
        """)
        
        st.markdown("---")
        st.markdown("### 📖 References")
        st.markdown("""
        - AASHTO LRFD Bridge Design
        - US Army Corps Manual
        - Thai Highway Standards
        """)
    
    # Main content tabs
    tab1, tab2, tab3, tab4 = st.tabs([
        "🏗️ Slope Geometry", 
        "🧱 Soil Layers", 
        "📊 Stability Analysis",
        "📉 Settlement Analysis"
    ])
    
    # ============================================================
    # Tab 1: Slope Geometry
    # ============================================================
    with tab1:
        st.markdown('<div class="sub-header">📐 Slope Configuration</div>', unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Embankment Dimensions**")
            
            # Check for loaded values
            default_height = float(st.session_state.get('loaded_slope', {}).get('height', 8.0))
            default_ratio = float(st.session_state.get('loaded_slope', {}).get('slope_ratio', 1.5))
            default_crest = float(st.session_state.get('loaded_slope', {}).get('crest_width', 10.0))
            default_toe = float(st.session_state.get('loaded_slope', {}).get('toe_x', 5.0))
            
            slope_height = st.number_input("Embankment Height (m)", 
                                           min_value=1.0, max_value=50.0, 
                                           value=default_height, step=0.5,
                                           help="ความสูงของคันทาง")
            
            slope_ratio = st.number_input("Slope Ratio (H:V)", 
                                          min_value=0.5, max_value=5.0, 
                                          value=default_ratio, step=0.1,
                                          help="อัตราส่วนความลาด (แนวนอน:แนวตั้ง)")
            
            crest_width = st.number_input("Crest Width (m)", 
                                          min_value=1.0, max_value=50.0, 
                                          value=default_crest, step=1.0,
                                          help="ความกว้างสันคันทาง")
            
            toe_x = st.number_input("Toe Position X (m)", 
                                    min_value=0.0, max_value=20.0, 
                                    value=default_toe, step=1.0,
                                    help="ตำแหน่ง X ของ toe")
            
            default_toe_elev = float(st.session_state.get('loaded_slope', {}).get('toe_elevation', 0.0))
            toe_elevation = st.number_input("Toe Elevation (m)", 
                                            min_value=-20.0, max_value=10.0, 
                                            value=default_toe_elev, step=0.5,
                                            help="ระดับความสูงของ toe (ค่าลบ = ต่ำกว่าระดับอ้างอิง)")
        
        with col2:
            st.markdown("**Water Table & Loading**")
            
            default_gwl = st.session_state.get('loaded_gwl', slope_height * 0.3 + toe_elevation)
            
            gwl = st.number_input("Ground Water Level (m)", 
                                  min_value=-25.0, max_value=float(slope_height + toe_elevation + 5), 
                                  value=float(default_gwl), step=0.5,
                                  help="ระดับน้ำใต้ดิน (Elevation)")
            
            surcharge = st.number_input("Surcharge Load (kPa)", 
                                        min_value=0.0, max_value=200.0, 
                                        value=10.0, step=5.0,
                                        help="น้ำหนักบรรทุกจร")
            
            st.markdown("---")
            st.markdown("**Calculated Values:**")
            slope_width = slope_height * slope_ratio
            slope_angle = np.degrees(np.arctan(1/slope_ratio))
            st.write(f"• Slope Width: **{slope_width:.2f} m**")
            st.write(f"• Slope Angle: **{slope_angle:.1f}°**")
            st.write(f"• Total Width: **{slope_width + crest_width:.2f} m**")
            st.write(f"• Crest Elevation: **{toe_elevation + slope_height:.2f} m**")
        
        # Store geometry
        slope_geometry = {
            'height': slope_height,
            'slope_ratio': slope_ratio,
            'crest_width': crest_width,
            'toe_x': toe_x,
            'toe_elevation': toe_elevation
        }
        st.session_state.slope_geometry = slope_geometry
        st.session_state.gwl = gwl
        st.session_state.surcharge = surcharge
        
        # Preview
        st.markdown('<div class="sub-header">🖼️ Geometry Preview</div>', unsafe_allow_html=True)
        
        fig_preview, ax_preview = plt.subplots(figsize=(12, 6))
        
        # Draw slope preview with toe_elevation
        slope_width = slope_height * slope_ratio
        crest_elev = toe_elevation + slope_height
        
        slope_x = [toe_x - 5, toe_x, toe_x + slope_width, 
                   toe_x + slope_width + crest_width, 
                   toe_x + slope_width + crest_width + 5]
        slope_y = [toe_elevation, toe_elevation, crest_elev, crest_elev, crest_elev]
        
        ax_preview.fill_between(slope_x, slope_y, toe_elevation - 3, color='#8B7355', alpha=0.7)
        ax_preview.plot(slope_x, slope_y, 'k-', linewidth=2)
        
        # Water table
        ax_preview.axhline(y=gwl, color='blue', linestyle='-', linewidth=2, label=f'GWL = {gwl:.1f} m')
        ax_preview.fill_between([toe_x - 5, toe_x + slope_width + crest_width + 5], 
                                toe_elevation - 3, gwl, color='lightblue', alpha=0.3)
        
        # Annotations
        ax_preview.annotate(f'H = {slope_height:.1f} m', 
                           xy=(toe_x - 2, toe_elevation + slope_height/2), fontsize=11, fontweight='bold')
        ax_preview.annotate(f'{slope_ratio}:1', 
                           xy=(toe_x + slope_width/2, toe_elevation + slope_height/2), fontsize=11, 
                           rotation=slope_angle, fontweight='bold', color='red')
        ax_preview.annotate(f'Crest = {crest_width:.1f} m',
                           xy=(toe_x + slope_width + crest_width/2, crest_elev + 0.5),
                           ha='center', fontsize=10)
        ax_preview.annotate(f'Toe Elev = {toe_elevation:.1f} m',
                           xy=(toe_x + 0.5, toe_elevation - 0.5),
                           ha='left', fontsize=10, color='brown')
        
        ax_preview.set_xlim(toe_x - 8, toe_x + slope_width + crest_width + 8)
        ax_preview.set_ylim(toe_elevation - 5, crest_elev + 3)
        ax_preview.set_xlabel('Distance (m)')
        ax_preview.set_ylabel('Elevation (m)')
        ax_preview.set_title('Embankment Geometry Preview')
        ax_preview.legend(loc='upper left')
        ax_preview.grid(True, alpha=0.3)
        ax_preview.set_aspect('equal')
        
        st.pyplot(fig_preview)
        plt.close()
    
    # ============================================================
    # Tab 2: Soil Layers
    # ============================================================
    with tab2:
        st.markdown('<div class="sub-header">🧱 Soil Layer Configuration (Max 8 Layers)</div>', unsafe_allow_html=True)
        
        # Number of layers
        n_layers = st.slider("Number of Soil Layers", 1, 8, len(st.session_state.soil_layers))
        
        # Adjust layer list
        while len(st.session_state.soil_layers) < n_layers:
            i = len(st.session_state.soil_layers) + 1
            st.session_state.soil_layers.append(
                SoilLayer(f"Layer {i}", 3.0, 18.0, 19.5, 20.0, 15.0, 15000, 0.3, 0.06, 0.75, 1.5, 1.5)
            )
        while len(st.session_state.soil_layers) > n_layers:
            st.session_state.soil_layers.pop()
        
        st.markdown("---")
        
        # Layer input
        for i in range(n_layers):
            with st.expander(f"📋 Layer {i+1}: {st.session_state.soil_layers[i].name}", expanded=(i == 0)):
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.markdown("**Basic Properties**")
                    name = st.text_input(f"Layer Name##{i}", 
                                        value=st.session_state.soil_layers[i].name,
                                        key=f"name_{i}")
                    thickness = st.number_input(f"Thickness (m)##{i}", 
                                               min_value=0.5, max_value=30.0,
                                               value=float(st.session_state.soil_layers[i].thickness),
                                               step=0.5, key=f"thick_{i}")
                    gamma = st.number_input(f"Unit Weight γ (kN/m³)##{i}", 
                                           min_value=10.0, max_value=25.0,
                                           value=float(st.session_state.soil_layers[i].gamma),
                                           step=0.5, key=f"gamma_{i}")
                    gamma_sat = st.number_input(f"Saturated Unit Weight γ_sat (kN/m³)##{i}", 
                                               min_value=15.0, max_value=28.0,
                                               value=float(st.session_state.soil_layers[i].gamma_sat),
                                               step=0.5, key=f"gamma_sat_{i}")
                
                with col2:
                    st.markdown("**Strength Parameters**")
                    cohesion = st.number_input(f"Cohesion c' (kPa)##{i}", 
                                              min_value=0.0, max_value=200.0,
                                              value=float(st.session_state.soil_layers[i].cohesion),
                                              step=1.0, key=f"c_{i}")
                    phi = st.number_input(f"Friction Angle φ' (°)##{i}", 
                                         min_value=0.0, max_value=45.0,
                                         value=float(st.session_state.soil_layers[i].phi),
                                         step=1.0, key=f"phi_{i}")
                    E = st.number_input(f"Young's Modulus E (kPa)##{i}", 
                                       min_value=1000.0, max_value=500000.0,
                                       value=float(st.session_state.soil_layers[i].E),
                                       step=1000.0, key=f"E_{i}")
                
                with col3:
                    st.markdown("**Consolidation Parameters**")
                    Cc = st.number_input(f"Compression Index Cc##{i}", 
                                        min_value=0.05, max_value=2.0,
                                        value=float(st.session_state.soil_layers[i].Cc),
                                        step=0.05, key=f"Cc_{i}")
                    Cr = st.number_input(f"Recompression Index Cr##{i}", 
                                        min_value=0.01, max_value=0.5,
                                        value=float(st.session_state.soil_layers[i].Cr),
                                        step=0.01, key=f"Cr_{i}")
                    e0 = st.number_input(f"Initial Void Ratio e₀##{i}", 
                                        min_value=0.3, max_value=3.0,
                                        value=float(st.session_state.soil_layers[i].e0),
                                        step=0.05, key=f"e0_{i}")
                    OCR = st.number_input(f"OCR##{i}", 
                                         min_value=1.0, max_value=10.0,
                                         value=float(st.session_state.soil_layers[i].OCR),
                                         step=0.5, key=f"OCR_{i}")
                    Cv = st.number_input(f"Cv (m²/year)##{i}", 
                                        min_value=0.1, max_value=50.0,
                                        value=float(st.session_state.soil_layers[i].Cv),
                                        step=0.5, key=f"Cv_{i}")
                
                # Update layer
                st.session_state.soil_layers[i] = SoilLayer(
                    name, thickness, gamma, gamma_sat, cohesion, phi, E, Cc, Cr, e0, OCR, Cv
                )
        
        # Summary table
        st.markdown('<div class="sub-header">📊 Soil Layers Summary</div>', unsafe_allow_html=True)
        
        summary_data = []
        total_thickness = 0
        for layer in st.session_state.soil_layers:
            summary_data.append({
                'Layer': layer.name,
                'Thickness (m)': layer.thickness,
                'γ (kN/m³)': layer.gamma,
                'γ_sat (kN/m³)': layer.gamma_sat,
                'c\' (kPa)': layer.cohesion,
                'φ\' (°)': layer.phi,
                'E (kPa)': layer.E,
                'Cc': layer.Cc,
                'e₀': layer.e0,
                'OCR': layer.OCR
            })
            total_thickness += layer.thickness
        
        st.dataframe(summary_data, use_container_width=True)
        st.info(f"📏 Total Soil Profile Thickness: **{total_thickness:.1f} m**")
        
        # Download configuration
        st.markdown("---")
        col_dl1, col_dl2 = st.columns(2)
        
        with col_dl1:
            json_data = save_to_json(
                st.session_state.get('slope_geometry', {'height': 8, 'slope_ratio': 1.5, 'crest_width': 10, 'toe_x': 5}),
                st.session_state.soil_layers,
                st.session_state.get('gwl', 2.0),
                {'surcharge': st.session_state.get('surcharge', 10)}
            )
            st.download_button(
                "📥 Download Configuration (JSON)",
                json_data,
                file_name="slope_config.json",
                mime="application/json"
            )
    
    # ============================================================
    # Tab 3: Stability Analysis
    # ============================================================
    with tab3:
        st.markdown('<div class="sub-header">📊 Slope Stability Analysis</div>', unsafe_allow_html=True)
        
        col1, col2 = st.columns([1, 2])
        
        with col1:
            st.markdown("**Analysis Settings**")
            
            analysis_method = st.selectbox(
                "Analysis Method",
                ["Bishop's Simplified", "Swedish (Ordinary Method)", "Both Methods"],
                help="เลือกวิธีการวิเคราะห์"
            )
            
            n_search = st.slider("Search Grid Density", 20, 100, 50,
                                help="จำนวน circle ในการค้นหา critical surface")
            
            show_slices = st.checkbox("Show Slice Division", value=True)
            
            st.markdown("---")
            st.markdown("**🌍 Seismic Analysis (Pseudo-static)**")
            
            enable_seismic = st.checkbox("Enable Seismic Analysis", value=False,
                                        help="เปิดใช้การวิเคราะห์แผ่นดินไหวแบบ Pseudo-static")
            
            if enable_seismic:
                seismic_zone = st.selectbox(
                    "Seismic Zone (Thailand)",
                    ["Zone 0 (kh=0.00)", "Zone 1 (kh=0.05)", "Zone 2 (kh=0.10)", 
                     "Zone 2A (kh=0.15)", "Zone 3 (kh=0.20)", "Custom"],
                    help="เลือกเขตแผ่นดินไหวตามมาตรฐาน มยผ."
                )
                
                if seismic_zone == "Custom":
                    seismic_coef = st.number_input("Horizontal Seismic Coefficient (kh)",
                                                   min_value=0.0, max_value=0.5,
                                                   value=0.10, step=0.01,
                                                   help="สัมประสิทธิ์แผ่นดินไหวแนวนอน")
                else:
                    # Extract kh from zone name
                    kh_map = {"Zone 0 (kh=0.00)": 0.0, "Zone 1 (kh=0.05)": 0.05, 
                              "Zone 2 (kh=0.10)": 0.10, "Zone 2A (kh=0.15)": 0.15,
                              "Zone 3 (kh=0.20)": 0.20}
                    seismic_coef = kh_map.get(seismic_zone, 0.10)
                    st.info(f"kh = {seismic_coef}")
            else:
                seismic_coef = 0.0
            
            st.session_state.seismic_coef = seismic_coef
            
            st.markdown("---")
            st.markdown("**Factor of Safety Criteria:**")
            st.write("• FS ≥ 1.50: ✅ Safe")
            st.write("• 1.00 ≤ FS < 1.50: ⚠️ Marginal")
            st.write("• FS < 1.00: ❌ Unsafe")
            
            if enable_seismic:
                st.markdown("**Seismic FS Criteria:**")
                st.write("• FS ≥ 1.10: ✅ Acceptable")
                st.write("• FS < 1.10: ❌ Needs improvement")
            
            run_analysis = st.button("🔄 Run Stability Analysis", type="primary", use_container_width=True)
        
        with col2:
            if run_analysis:
                slope_geometry = st.session_state.get('slope_geometry')
                gwl = st.session_state.get('gwl', 2.0)
                
                if slope_geometry and st.session_state.soil_layers:
                    with st.spinner("Searching for critical slip circle..."):
                        
                        if analysis_method == "Both Methods":
                            # Run both analyses
                            result_bishop = search_critical_circle(
                                slope_geometry, st.session_state.soil_layers,
                                gwl, "Bishop", n_search, seismic_coef
                            )
                            result_swedish = search_critical_circle(
                                slope_geometry, st.session_state.soil_layers,
                                gwl, "Swedish", n_search, seismic_coef
                            )
                            
                            # Display comparison
                            st.markdown("### 📈 Results Comparison")
                            
                            col_b, col_s = st.columns(2)
                            
                            with col_b:
                                if result_bishop:
                                    fs_b = result_bishop.fs
                                    status = 'safe' if fs_b >= 1.5 else 'warning' if fs_b >= 1.0 else 'danger'
                                    st.markdown(f"""
                                    <div class="result-{status}">
                                        <h3>Bishop's Simplified</h3>
                                        <h1>FS = {fs_b:.3f}</h1>
                                        <p>Iterations: {result_bishop.iterations}</p>
                                    </div>
                                    """, unsafe_allow_html=True)
                            
                            with col_s:
                                if result_swedish:
                                    fs_s = result_swedish.fs
                                    status = 'safe' if fs_s >= 1.5 else 'warning' if fs_s >= 1.0 else 'danger'
                                    st.markdown(f"""
                                    <div class="result-{status}">
                                        <h3>Swedish Method</h3>
                                        <h1>FS = {fs_s:.3f}</h1>
                                        <p>No iteration</p>
                                    </div>
                                    """, unsafe_allow_html=True)
                            
                            # Use Bishop for plotting (more accurate)
                            result = result_bishop
                            st.session_state.analysis_result = result
                            
                        else:
                            method = "Bishop" if "Bishop" in analysis_method else "Swedish"
                            result = search_critical_circle(
                                slope_geometry, st.session_state.soil_layers,
                                gwl, method, n_search, seismic_coef
                            )
                            st.session_state.analysis_result = result
                            
                            if result:
                                fs = result.fs
                                status = 'safe' if fs >= 1.5 else 'warning' if fs >= 1.0 else 'danger'
                                status_text = 'SAFE' if fs >= 1.5 else 'MARGINAL' if fs >= 1.0 else 'UNSAFE'
                                
                                st.markdown(f"""
                                <div class="result-{status}">
                                    <h2>{result.method}</h2>
                                    <h1 style="font-size:3rem;">FS = {fs:.3f}</h1>
                                    <h3>{status_text}</h3>
                                </div>
                                """, unsafe_allow_html=True)
                else:
                    st.error("Please configure slope geometry and soil layers first!")
        
        # Plot results
        if st.session_state.analysis_result:
            st.markdown('<div class="sub-header">🖼️ Analysis Results Visualization</div>', unsafe_allow_html=True)
            
            result = st.session_state.analysis_result
            slope_geometry = st.session_state.slope_geometry
            gwl = st.session_state.gwl
            
            fig = plot_slope_and_circle(
                slope_geometry, 
                st.session_state.soil_layers,
                gwl, 
                result,
                show_slices
            )
            st.pyplot(fig)
            plt.close()
            
            # Detailed slice data
            with st.expander("📋 Detailed Slice Data"):
                if result.slices_data:
                    slice_df = []
                    for s in result.slices_data:
                        slice_df.append({
                            'Slice': s['index'] + 1,
                            'x_mid (m)': f"{s['x_mid']:.2f}",
                            'Width (m)': f"{s['width']:.2f}",
                            'Height (m)': f"{s['height']:.2f}",
                            'W (kN)': f"{s['W']:.1f}",
                            'α (°)': f"{s['alpha_deg']:.1f}",
                            'u (kPa)': f"{s.get('u', 0):.1f}",
                            'c\' (kPa)': f"{s['c']:.1f}",
                            'φ\' (°)': f"{s['phi']:.1f}",
                            'Soil': s['soil_name']
                        })
                    st.dataframe(slice_df, use_container_width=True)
            
            # Critical circle info
            with st.expander("🎯 Critical Circle Parameters"):
                circle = result.critical_circle
                st.write(f"• Center X: **{circle.x_center:.2f} m**")
                st.write(f"• Center Y: **{circle.y_center:.2f} m**")
                st.write(f"• Radius: **{circle.radius:.2f} m**")
            
            # Export to Word
            st.markdown("---")
            st.markdown("### 📄 Export Report")
            
            col_exp1, col_exp2 = st.columns(2)
            
            with col_exp1:
                try:
                    # Generate figure for report
                    fig_for_report = plot_slope_and_circle(
                        slope_geometry, 
                        st.session_state.soil_layers,
                        gwl, 
                        result,
                        True
                    )
                    
                    seismic_coef = st.session_state.get('seismic_coef', 0.0)
                    
                    word_bytes = generate_word_report(
                        slope_geometry,
                        st.session_state.soil_layers,
                        gwl,
                        result,
                        seismic_coef,
                        fig_for_report
                    )
                    plt.close(fig_for_report)
                    
                    st.download_button(
                        label="📥 Download Word Report (.docx)",
                        data=word_bytes,
                        file_name="slope_stability_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                except ImportError:
                    st.warning("python-docx library not installed. Install with: pip install python-docx")
                except Exception as e:
                    st.error(f"Error generating report: {str(e)}")
            
            with col_exp2:
                # Save figure as PNG
                fig_png = plot_slope_and_circle(
                    slope_geometry, 
                    st.session_state.soil_layers,
                    gwl, 
                    result,
                    True
                )
                
                buf = BytesIO()
                fig_png.savefig(buf, format='png', dpi=200, bbox_inches='tight')
                buf.seek(0)
                plt.close(fig_png)
                
                st.download_button(
                    label="📥 Download Figure (.png)",
                    data=buf.getvalue(),
                    file_name="slope_stability_figure.png",
                    mime="image/png",
                    use_container_width=True
                )
    
    # ============================================================
    # Tab 4: Settlement Analysis
    # ============================================================
    with tab4:
        st.markdown('<div class="sub-header">📉 Settlement Analysis</div>', unsafe_allow_html=True)
        
        col1, col2 = st.columns([1, 2])
        
        with col1:
            st.markdown("**Loading Parameters**")
            
            # Applied stress
            q_applied = st.number_input("Applied Stress q (kPa)",
                                        min_value=10.0, max_value=500.0,
                                        value=float(st.session_state.get('surcharge', 50.0)),
                                        step=10.0,
                                        help="น้ำหนักบรรทุกที่กระทำ")
            
            foundation_width = st.number_input("Foundation Width B (m)",
                                               min_value=1.0, max_value=50.0,
                                               value=10.0, step=1.0,
                                               help="ความกว้างฐานราก/คันทาง")
            
            foundation_depth = st.number_input("Foundation Depth (m)",
                                               min_value=0.0, max_value=10.0,
                                               value=0.0, step=0.5)
            
            st.markdown("---")
            st.markdown("**Time Parameters**")
            
            analysis_time = st.number_input("Analysis Period (years)",
                                            min_value=1.0, max_value=100.0,
                                            value=30.0, step=5.0)
            
            drainage_condition = st.selectbox(
                "Drainage Condition",
                ["Single Drainage (Top)", "Double Drainage (Top & Bottom)"],
                help="เงื่อนไขการระบายน้ำ"
            )
            
            run_settlement = st.button("🔄 Run Settlement Analysis", 
                                       type="primary", use_container_width=True)
        
        with col2:
            if run_settlement and st.session_state.soil_layers:
                with st.spinner("Calculating settlement..."):
                    # Immediate settlement
                    Si = calculate_immediate_settlement(
                        q_applied, foundation_width, 
                        st.session_state.soil_layers
                    )
                    
                    # Consolidation settlement
                    Sc, layer_settlements = calculate_consolidation_settlement(
                        q_applied, st.session_state.soil_layers, foundation_depth
                    )
                    
                    # Total ultimate settlement
                    S_total = Si + Sc
                    
                    # Time-settlement curve
                    total_thickness = sum(l.thickness for l in st.session_state.soil_layers)
                    H_drainage = total_thickness if "Single" in drainage_condition else total_thickness / 2
                    avg_Cv = np.mean([l.Cv for l in st.session_state.soil_layers])
                    
                    times = np.linspace(0, analysis_time, 50)
                    settlement_time = calculate_time_rate_settlement(Sc, avg_Cv, H_drainage, times)
                    
                    # Display results
                    st.markdown("### 📊 Settlement Results")
                    
                    col_r1, col_r2, col_r3 = st.columns(3)
                    
                    with col_r1:
                        st.metric("Immediate Settlement (Si)", f"{Si:.1f} mm")
                    with col_r2:
                        st.metric("Consolidation Settlement (Sc)", f"{Sc:.1f} mm")
                    with col_r3:
                        st.metric("Total Settlement (S)", f"{S_total:.1f} mm")
                    
                    # Settlement by layer
                    st.markdown("---")
                    st.markdown("**Settlement by Layer:**")
                    
                    layer_df = []
                    for ls in layer_settlements:
                        layer_df.append({
                            'Layer': ls['layer'],
                            'Thickness (m)': ls['thickness'],
                            'σ\'₀ (kPa)': f"{ls['sigma_v0']:.1f}",
                            'Δσ (kPa)': f"{ls['delta_sigma']:.1f}",
                            'Settlement (mm)': f"{ls['settlement']:.2f}"
                        })
                    st.dataframe(layer_df, use_container_width=True)
                    
                    # Store for plotting
                    st.session_state.settlement_data = {
                        'Si': Si,
                        'Sc': Sc,
                        'S_total': S_total,
                        'time_data': settlement_time,
                        'layer_data': layer_settlements
                    }
        
        # Plots
        if 'settlement_data' in st.session_state:
            st.markdown('<div class="sub-header">🖼️ Settlement Visualization</div>', unsafe_allow_html=True)
            
            tab_s1, tab_s2 = st.tabs(["📈 Time-Settlement Curve", "📊 Stress Distribution"])
            
            with tab_s1:
                data = st.session_state.settlement_data
                fig_time = plot_settlement_time(data['time_data'], data['Sc'])
                st.pyplot(fig_time)
                plt.close()
                
                # Time to reach specific consolidation
                st.markdown("**Consolidation Time Estimates:**")
                for pct in [50, 90, 95]:
                    target = data['Sc'] * pct / 100
                    for t, s in data['time_data']:
                        if s >= target:
                            st.write(f"• {pct}% Consolidation: **{t:.1f} years** ({target:.1f} mm)")
                            break
            
            with tab_s2:
                fig_stress = plot_stress_distribution(
                    st.session_state.soil_layers,
                    q_applied
                )
                st.pyplot(fig_stress)
                plt.close()
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align:center; color:#666; padding:1rem;">
        <p><strong>Slope Stability & Settlement Analysis</strong></p>
        <p>วิเคราะห์เสถียรภาพลาดดินคันทางด้วยวิธี Bishop และ Swedish</p>
        <p>Developed for Civil Engineering Education | KMUTNB</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
