import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib as mpl
import json
from io import BytesIO

# ตั้งค่าฟอนต์ไทยสำหรับ matplotlib
try:
    import matplotlib.font_manager as fm
    # ใช้ฟอนต์ Loma ที่รองรับภาษาไทย
    plt.rcParams['font.family'] = 'Loma'
    # ตั้งค่าให้แสดงเครื่องหมายลบได้ถูกต้อง
    plt.rcParams['axes.unicode_minus'] = False
except:
    pass

# =========================================================
# Functions
# =========================================================
def concrete_mix_design(
    wc_ratio,
    max_agg_mm,
    sg_cement,
    sg_fine,
    sg_coarse,
    air_content,
    unit_weight_coarse
):
    """
    ACI 211.1 Concrete Mix Design Method
    Returns: dict with Water, Cement, Fine Aggregate, Coarse Aggregate (kg/m³)
    """
    # ---- Water content & coarse aggregate volume (ACI typical) ----
    if max_agg_mm == 20:
        water = 185
        vol_coarse = 0.62
    elif max_agg_mm == 25:
        water = 175
        vol_coarse = 0.64
    else:  # 40 mm
        water = 165
        vol_coarse = 0.68

    cement = water / wc_ratio
    weight_coarse = vol_coarse * unit_weight_coarse

    # ---- Volume calculations ----
    vol_water = water / 1000
    vol_cement = cement / (sg_cement * 1000)
    vol_coarse_abs = weight_coarse / (sg_coarse * 1000)

    vol_fine = 1 - (
        vol_water +
        vol_cement +
        vol_coarse_abs +
        air_content
    )

    weight_fine = vol_fine * sg_fine * 1000

    return {
        "Water": water,
        "Cement": cement,
        "Fine Aggregate": weight_fine,
        "Coarse Aggregate": weight_coarse,
        # เก็บค่ากลางสำหรับรายงาน
        "vol_water": vol_water,
        "vol_cement": vol_cement,
        "vol_coarse": vol_coarse_abs,
        "vol_fine": vol_fine,
        "vol_air": air_content
    }


def moisture_correction(weight_ssd, mc, absorption):
    """
    Moisture correction for aggregates
    weight_ssd : SSD weight (kg/m³)
    mc         : moisture content (%)
    absorption : absorption (%)
    Returns: (delta_water, batch_weight)
    """
    delta_water = weight_ssd * (mc - absorption) / 100
    batch_weight = weight_ssd * (1 + mc / 100)
    return delta_water, batch_weight


def create_word_report(input_data, mix_result, moisture_result):
    """
    สร้างรายงาน Word แบบละเอียดเป็นขั้นเป็นตอน
    ใช้ python-docx
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        import subprocess
        subprocess.run(['pip', 'install', 'python-docx', '--break-system-packages'], 
                      capture_output=True)
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    
    # สร้างเอกสาร
    doc = Document()
    
    # ฟังก์ชันตั้งค่าฟอนต์ไทย
    def set_thai_font(run, size=15, bold=False):
        run.font.name = 'TH SarabunPSK'
        run.font.size = Pt(size)
        run.font.bold = bold
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
    
    # หัวเรื่อง
    title = doc.add_heading('รายงานการออกแบบส่วนผสมคอนกรีต', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_thai_font(run, size=18, bold=True)
    
    subtitle = doc.add_paragraph('ตามมาตรฐาน ACI 211.1')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_thai_font(run, size=14)
    
    # ส่วนที่ 1: ข้อมูลนำเข้า
    doc.add_heading('1. ข้อมูลนำเข้าในการออกแบบ', 1)
    
    table1 = doc.add_table(rows=9, cols=2)
    table1.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'รายการ'
    hdr_cells[1].text = 'ค่าที่ใช้'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_thai_font(run, bold=True)
    
    # Data
    data1 = [
        ('อัตราส่วน น้ำ/ปูนซีเมนต์ (w/c)', f"{input_data['wc_ratio']}"),
        ('ขนาดมวลรวมหยาบสูงสุด (mm)', f"{input_data['max_agg_mm']}"),
        ('ค่าความถ่วงจำเพาะของปูนซีเมนต์', f"{input_data['sg_cement']:.2f}"),
        ('ค่าความถ่วงจำเพาะของมวลรวมละเอียด', f"{input_data['sg_fine']:.2f}"),
        ('ค่าความถ่วงจำเพาะของมวลรวมหยาบ', f"{input_data['sg_coarse']:.2f}"),
        ('ปริมาณอากาศ (%)', f"{input_data['air_content']*100:.1f}"),
        ('น้ำหนักหน่วยของมวลรวมหยาบ (kg/m³)', f"{input_data['unit_weight_coarse']:.0f}")
    ]
    
    for i, (label, value) in enumerate(data1, start=1):
        row_cells = table1.rows[i].cells
        row_cells[0].text = label
        row_cells[1].text = value
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_thai_font(run)
    
    # ส่วนที่ 2: ขั้นตอนการคำนวณ
    doc.add_heading('2. ขั้นตอนการคำนวณตามวิธี ACI 211.1', 1)
    
    # ขั้นตอนที่ 1
    p = doc.add_paragraph()
    run = p.add_run('ขั้นตอนที่ 1: กำหนดปริมาณน้ำและปริมาณมวลรวมหยาบ')
    set_thai_font(run, bold=True)
    
    vol_coarse_ratio = 0.62 if input_data['max_agg_mm'] == 20 else (0.64 if input_data['max_agg_mm'] == 25 else 0.68)
    text1 = f"จากตาราง ACI สำหรับขนาดมวลรวมหยาบสูงสุด {input_data['max_agg_mm']} mm:\n"
    text1 += f"  - ปริมาณน้ำ = {mix_result['Water']:.1f} kg/m³\n"
    text1 += f"  - สัดส่วนปริมาตรมวลรวมหยาบ = {vol_coarse_ratio}"
    p = doc.add_paragraph(text1)
    for run in p.runs:
        set_thai_font(run)
    
    # ขั้นตอนที่ 2
    p = doc.add_paragraph()
    run = p.add_run('ขั้นตอนที่ 2: คำนวณปริมาณปูนซีเมนต์')
    set_thai_font(run, bold=True)
    
    text2 = f"ปริมาณปูนซีเมนต์ = น้ำ / (w/c) = {mix_result['Water']:.1f} / {input_data['wc_ratio']} = {mix_result['Cement']:.1f} kg/m³"
    p = doc.add_paragraph(text2)
    for run in p.runs:
        set_thai_font(run)
    
    # ขั้นตอนที่ 3
    p = doc.add_paragraph()
    run = p.add_run('ขั้นตอนที่ 3: คำนวณน้ำหนักมวลรวมหยาบ')
    set_thai_font(run, bold=True)
    
    text3 = f"น้ำหนักมวลรวมหยาบ = สัดส่วนปริมาตร × น้ำหนักหน่วย\n"
    text3 += f"  = {vol_coarse_ratio} × {input_data['unit_weight_coarse']} = {mix_result['Coarse Aggregate']:.1f} kg/m³"
    p = doc.add_paragraph(text3)
    for run in p.runs:
        set_thai_font(run)
    
    # ขั้นตอนที่ 4
    p = doc.add_paragraph()
    run = p.add_run('ขั้นตอนที่ 4: คำนวณปริมาตรสัมบูรณ์ของแต่ละวัสดุ')
    set_thai_font(run, bold=True)
    
    text4 = f"ปริมาตรน้ำ = {mix_result['Water']:.1f} / 1000 = {mix_result['vol_water']:.4f} m³\n"
    text4 += f"ปริมาตรปูนซีเมนต์ = {mix_result['Cement']:.1f} / ({input_data['sg_cement']} × 1000) = {mix_result['vol_cement']:.4f} m³\n"
    text4 += f"ปริมาตรมวลรวมหยาบ = {mix_result['Coarse Aggregate']:.1f} / ({input_data['sg_coarse']} × 1000) = {mix_result['vol_coarse']:.4f} m³\n"
    text4 += f"ปริมาตรอากาศ = {input_data['air_content']*100:.1f}% = {mix_result['vol_air']:.4f} m³"
    p = doc.add_paragraph(text4)
    for run in p.runs:
        set_thai_font(run)
    
    # ขั้นตอนที่ 5
    p = doc.add_paragraph()
    run = p.add_run('ขั้นตอนที่ 5: คำนวณปริมาตรมวลรวมละเอียด')
    set_thai_font(run, bold=True)
    
    text5 = f"ปริมาตรมวลรวมละเอียด = 1 - (น้ำ + ปูนซีเมนต์ + มวลรวมหยาบ + อากาศ)\n"
    text5 += f"  = 1 - ({mix_result['vol_water']:.4f} + {mix_result['vol_cement']:.4f} + {mix_result['vol_coarse']:.4f} + {mix_result['vol_air']:.4f})\n"
    text5 += f"  = {mix_result['vol_fine']:.4f} m³"
    p = doc.add_paragraph(text5)
    for run in p.runs:
        set_thai_font(run)
    
    # ขั้นตอนที่ 6
    p = doc.add_paragraph()
    run = p.add_run('ขั้นตอนที่ 6: คำนวณน้ำหนักมวลรวมละเอียด')
    set_thai_font(run, bold=True)
    
    text6 = f"น้ำหนักมวลรวมละเอียด = ปริมาตร × ความถ่วงจำเพาะ × 1000\n"
    text6 += f"  = {mix_result['vol_fine']:.4f} × {input_data['sg_fine']} × 1000\n"
    text6 += f"  = {mix_result['Fine Aggregate']:.1f} kg/m³"
    p = doc.add_paragraph(text6)
    for run in p.runs:
        set_thai_font(run)
    
    # ส่วนที่ 3: ผลลัพธ์ SSD
    doc.add_heading('3. ผลการออกแบบส่วนผสมคอนกรีต (สภาพ SSD)', 1)
    
    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Light Grid Accent 1'
    
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'วัสดุ'
    hdr_cells[1].text = 'ปริมาณ (kg/m³)'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_thai_font(run, bold=True)
    
    data2 = [
        ('น้ำ', f"{mix_result['Water']:.1f}"),
        ('ปูนซีเมนต์', f"{mix_result['Cement']:.1f}"),
        ('มวลรวมละเอียด', f"{mix_result['Fine Aggregate']:.1f}"),
        ('มวลรวมหยาบ', f"{mix_result['Coarse Aggregate']:.1f}")
    ]
    
    for i, (label, value) in enumerate(data2, start=1):
        row_cells = table2.rows[i].cells
        row_cells[0].text = label
        row_cells[1].text = value
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_thai_font(run)
    
    # ส่วนที่ 4: การปรับแก้ความชื้น
    doc.add_heading('4. การปรับแก้เนื่องจากความชื้นในมวลรวม', 1)
    
    # 4.1 มวลรวมละเอียด
    p = doc.add_paragraph()
    run = p.add_run('4.1 การคำนวณสำหรับมวลรวมละเอียด')
    set_thai_font(run, bold=True)
    
    text_mc1 = f"ความชื้น (MC) = {input_data['mc_fine']:.1f}%\n"
    text_mc1 += f"การดูดซับน้ำ (Absorption) = {input_data['abs_fine']:.1f}%\n"
    text_mc1 += f"การเปลี่ยนแปลงน้ำหนักน้ำ = น้ำหนัก SSD × (MC - Absorption) / 100\n"
    text_mc1 += f"  = {mix_result['Fine Aggregate']:.1f} × ({input_data['mc_fine']:.1f} - {input_data['abs_fine']:.1f}) / 100\n"
    text_mc1 += f"  = {moisture_result['dw_fine']:.1f} kg/m³\n\n"
    text_mc1 += f"น้ำหนักมวลรวมละเอียดสำหรับผสม = น้ำหนัก SSD × (1 + MC/100)\n"
    text_mc1 += f"  = {mix_result['Fine Aggregate']:.1f} × (1 + {input_data['mc_fine']:.1f}/100)\n"
    text_mc1 += f"  = {moisture_result['batch_fine']:.1f} kg/m³"
    p = doc.add_paragraph(text_mc1)
    for run in p.runs:
        set_thai_font(run)
    
    # 4.2 มวลรวมหยาบ
    p = doc.add_paragraph()
    run = p.add_run('4.2 การคำนวณสำหรับมวลรวมหยาบ')
    set_thai_font(run, bold=True)
    
    text_mc2 = f"ความชื้น (MC) = {input_data['mc_coarse']:.1f}%\n"
    text_mc2 += f"การดูดซับน้ำ (Absorption) = {input_data['abs_coarse']:.1f}%\n"
    text_mc2 += f"การเปลี่ยนแปลงน้ำหนักน้ำ = น้ำหนัก SSD × (MC - Absorption) / 100\n"
    text_mc2 += f"  = {mix_result['Coarse Aggregate']:.1f} × ({input_data['mc_coarse']:.1f} - {input_data['abs_coarse']:.1f}) / 100\n"
    text_mc2 += f"  = {moisture_result['dw_coarse']:.1f} kg/m³\n\n"
    text_mc2 += f"น้ำหนักมวลรวมหยาบสำหรับผสม = น้ำหนัก SSD × (1 + MC/100)\n"
    text_mc2 += f"  = {mix_result['Coarse Aggregate']:.1f} × (1 + {input_data['mc_coarse']:.1f}/100)\n"
    text_mc2 += f"  = {moisture_result['batch_coarse']:.1f} kg/m³"
    p = doc.add_paragraph(text_mc2)
    for run in p.runs:
        set_thai_font(run)
    
    # 4.3 ปรับแก้น้ำ
    p = doc.add_paragraph()
    run = p.add_run('4.3 การปรับแก้ปริมาณน้ำผสม')
    set_thai_font(run, bold=True)
    
    text_mc3 = f"น้ำที่มาจากมวลรวมทั้งหมด = {moisture_result['dw_fine']:.1f} + {moisture_result['dw_coarse']:.1f} = {moisture_result['total_delta_water']:.1f} kg/m³\n"
    text_mc3 += f"ปริมาณน้ำผสมที่ต้องเติม = {mix_result['Water']:.1f} - {moisture_result['total_delta_water']:.1f} = {moisture_result['corrected_water']:.1f} kg/m³"
    p = doc.add_paragraph(text_mc3)
    for run in p.runs:
        set_thai_font(run)
    
    # ส่วนที่ 5: สรุปส่วนผสม
    doc.add_heading('5. สรุปส่วนผสมคอนกรีตสำหรับการผสม', 1)
    
    table3 = doc.add_table(rows=5, cols=3)
    table3.style = 'Light Grid Accent 1'
    
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = 'วัสดุ'
    hdr_cells[1].text = 'SSD (kg/m³)'
    hdr_cells[2].text = 'สำหรับผสม (kg/m³)'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_thai_font(run, bold=True)
    
    data3 = [
        ('น้ำผสม', f"{mix_result['Water']:.1f}", f"{moisture_result['corrected_water']:.1f}"),
        ('ปูนซีเมนต์', f"{mix_result['Cement']:.1f}", f"{mix_result['Cement']:.1f}"),
        ('มวลรวมละเอียด', f"{mix_result['Fine Aggregate']:.1f}", f"{moisture_result['batch_fine']:.1f}"),
        ('มวลรวมหยาบ', f"{mix_result['Coarse Aggregate']:.1f}", f"{moisture_result['batch_coarse']:.1f}")
    ]
    
    for i, (label, ssd, batch) in enumerate(data3, start=1):
        row_cells = table3.rows[i].cells
        row_cells[0].text = label
        row_cells[1].text = ssd
        row_cells[2].text = batch
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_thai_font(run)
    
    # หมายเหตุ
    p = doc.add_paragraph('หมายเหตุ: คอลัมน์ "สำหรับผสม" คือส่วนผสมที่ต้องใช้ในการผสมคอนกรีตจริง')
    for run in p.runs:
        set_thai_font(run)
        run.italic = True
    
    # บันทึกไฟล์
    output_path = '/home/claude/concrete_mix_report.docx'
    doc.save(output_path)
    
    return output_path


# =========================================================
# Streamlit UI
# =========================================================
st.set_page_config(page_title="Concrete Mix Design - ACI 211", layout="centered")

st.title("🏗️ โปรแกรมออกแบบส่วนผสมคอนกรีต")
st.caption("ตามมาตรฐาน ACI 211.1 | สำหรับการสอนและงานปฏิบัติ")

# ---------------- Upload JSON ----------------
st.sidebar.header("📂 โหลดข้อมูลจากไฟล์")

uploaded_json = st.sidebar.file_uploader("อัพโหลดไฟล์ JSON", type=['json'])

if uploaded_json is not None:
    try:
        loaded_data = json.load(uploaded_json)
        
        # ตรวจสอบว่าเป็นไฟล์ใหม่
        file_id = f"{uploaded_json.name}_{uploaded_json.size}"
        if st.session_state.get('last_uploaded_file') != file_id:
            st.session_state['last_uploaded_file'] = file_id
            
            # อัพเดท session_state
            st.session_state['input_wc_ratio'] = loaded_data.get('wc_ratio', 0.50)
            st.session_state['input_max_agg_mm'] = loaded_data.get('max_agg_mm', 25)
            st.session_state['input_sg_cement'] = loaded_data.get('sg_cement', 3.15)
            st.session_state['input_sg_fine'] = loaded_data.get('sg_fine', 2.65)
            st.session_state['input_sg_coarse'] = loaded_data.get('sg_coarse', 2.70)
            st.session_state['input_air_content'] = loaded_data.get('air_content', 2.0)
            st.session_state['input_unit_weight_coarse'] = loaded_data.get('unit_weight_coarse', 1600)
            st.session_state['input_mc_fine'] = loaded_data.get('mc_fine', 5.0)
            st.session_state['input_abs_fine'] = loaded_data.get('abs_fine', 2.0)
            st.session_state['input_mc_coarse'] = loaded_data.get('mc_coarse', 1.0)
            st.session_state['input_abs_coarse'] = loaded_data.get('abs_coarse', 0.5)
            
            st.sidebar.success("✅ โหลดข้อมูลสำเร็จ!")
            st.rerun()
            
    except Exception as e:
        st.sidebar.error(f"❌ ไม่สามารถอ่านไฟล์ได้: {e}")

# ---------------- Sidebar Inputs ----------------
st.sidebar.header("📥 ข้อมูลสำหรับออกแบบส่วนผสม")

wc_ratio = st.sidebar.number_input(
    "อัตราส่วน น้ำ/ปูนซีเมนต์ (w/c)",
    min_value=0.35, max_value=0.70, step=0.01,
    value=st.session_state.get('input_wc_ratio', 0.50),
    key="input_wc_ratio"
)

max_agg_options = [20, 25, 40]
current_max_agg = st.session_state.get('input_max_agg_mm', 25)
max_agg_idx = max_agg_options.index(current_max_agg) if current_max_agg in max_agg_options else 1

max_agg_mm = st.sidebar.selectbox(
    "ขนาดมวลรวมหยาบสูงสุด (mm)",
    options=max_agg_options,
    index=max_agg_idx,
    key="input_max_agg_mm"
)

sg_cement = st.sidebar.number_input(
    "ค่าความถ่วงจำเพาะของปูนซีเมนต์",
    value=st.session_state.get('input_sg_cement', 3.15),
    key="input_sg_cement"
)

sg_fine = st.sidebar.number_input(
    "ค่าความถ่วงจำเพาะของมวลรวมละเอียด",
    value=st.session_state.get('input_sg_fine', 2.65),
    key="input_sg_fine"
)

sg_coarse = st.sidebar.number_input(
    "ค่าความถ่วงจำเพาะของมวลรวมหยาบ",
    value=st.session_state.get('input_sg_coarse', 2.70),
    key="input_sg_coarse"
)

air_content = st.sidebar.slider(
    "ปริมาณอากาศ (%)",
    min_value=1.0, max_value=6.0,
    value=st.session_state.get('input_air_content', 2.0),
    key="input_air_content"
) / 100

unit_weight_coarse = st.sidebar.number_input(
    "น้ำหนักหน่วยของมวลรวมหยาบ (kg/m³)",
    value=st.session_state.get('input_unit_weight_coarse', 1600),
    key="input_unit_weight_coarse"
)

# ---------------- Moisture Input ----------------
st.sidebar.header("💧 ข้อมูลความชื้นของมวลรวม")

mc_fine = st.sidebar.number_input(
    "ความชื้นมวลรวมละเอียด (%)",
    value=st.session_state.get('input_mc_fine', 5.0),
    key="input_mc_fine"
)

abs_fine = st.sidebar.number_input(
    "การดูดซับน้ำมวลรวมละเอียด (%)",
    value=st.session_state.get('input_abs_fine', 2.0),
    key="input_abs_fine"
)

mc_coarse = st.sidebar.number_input(
    "ความชื้นมวลรวมหยาบ (%)",
    value=st.session_state.get('input_mc_coarse', 1.0),
    key="input_mc_coarse"
)

abs_coarse = st.sidebar.number_input(
    "การดูดซับน้ำมวลรวมหยาบ (%)",
    value=st.session_state.get('input_abs_coarse', 0.5),
    key="input_abs_coarse"
)

# ---------------- Download JSON Button ----------------
export_data = {
    'wc_ratio': wc_ratio,
    'max_agg_mm': max_agg_mm,
    'sg_cement': sg_cement,
    'sg_fine': sg_fine,
    'sg_coarse': sg_coarse,
    'air_content': air_content * 100,
    'unit_weight_coarse': unit_weight_coarse,
    'mc_fine': mc_fine,
    'abs_fine': abs_fine,
    'mc_coarse': mc_coarse,
    'abs_coarse': abs_coarse
}

json_str = json.dumps(export_data, ensure_ascii=False, indent=2)

st.sidebar.download_button(
    label="💾 ดาวน์โหลดข้อมูล (JSON)",
    data=json_str,
    file_name="concrete_mix_input.json",
    mime="application/json"
)

# =========================================================
# Calculation
# =========================================================
if st.button("🧮 คำนวณส่วนผสมคอนกรีต", type="primary"):

    # ---- Mix design ----
    mix = concrete_mix_design(
        wc_ratio,
        max_agg_mm,
        sg_cement,
        sg_fine,
        sg_coarse,
        air_content,
        unit_weight_coarse
    )

    df_mix = pd.DataFrame({
        "วัสดุ": ["น้ำ", "ปูนซีเมนต์", "มวลรวมละเอียด", "มวลรวมหยาบ"],
        "ปริมาณ (kg/m³)": [
            round(mix["Water"], 1),
            round(mix["Cement"], 1),
            round(mix["Fine Aggregate"], 1),
            round(mix["Coarse Aggregate"], 1)
        ]
    })

    st.subheader("📊 ผลการออกแบบส่วนผสม (สภาพ SSD)")
    st.dataframe(df_mix, use_container_width=True)

    # ---- Moisture correction ----
    dw_fine, batch_fine = moisture_correction(
        mix["Fine Aggregate"], mc_fine, abs_fine
    )

    dw_coarse, batch_coarse = moisture_correction(
        mix["Coarse Aggregate"], mc_coarse, abs_coarse
    )

    total_delta_water = dw_fine + dw_coarse
    corrected_water = mix["Water"] - total_delta_water

    df_mc = pd.DataFrame({
        "รายการ": ["มวลรวมละเอียด", "มวลรวมหยาบ"],
        "น้ำหนัก SSD (kg/m³)": [
            round(mix["Fine Aggregate"], 1),
            round(mix["Coarse Aggregate"], 1)
        ],
        "น้ำหนักสำหรับผสม (kg/m³)": [
            round(batch_fine, 1),
            round(batch_coarse, 1)
        ],
        "Δ น้ำ (kg/m³)": [
            round(dw_fine, 1),
            round(dw_coarse, 1)
        ]
    })

    st.subheader("💧 การปรับแก้เนื่องจากความชื้น")
    st.dataframe(df_mc, use_container_width=True)

    st.info(
        f"💧 **ปริมาณน้ำเดิม** = {round(mix['Water'],1)} kg/m³\n\n"
        f"💧 **ปริมาณน้ำผสมที่ปรับแก้** = {round(corrected_water,1)} kg/m³"
    )

    # ---- Summary table ----
    st.subheader("✅ สรุปส่วนผสมคอนกรีตสำหรับการผสม")
    
    df_final = pd.DataFrame({
        "วัสดุ": ["น้ำผสม", "ปูนซีเมนต์", "มวลรวมละเอียด", "มวลรวมหยาบ"],
        "SSD (kg/m³)": [
            round(mix["Water"], 1),
            round(mix["Cement"], 1),
            round(mix["Fine Aggregate"], 1),
            round(mix["Coarse Aggregate"], 1)
        ],
        "สำหรับผสม (kg/m³)": [
            round(corrected_water, 1),
            round(mix["Cement"], 1),
            round(batch_fine, 1),
            round(batch_coarse, 1)
        ]
    })
    
    st.dataframe(df_final, use_container_width=True)

    # ---- Pie Chart ----
    st.subheader("📈 สัดส่วนวัสดุ (สภาพ SSD)")
    
    # เตรียม labels ในภาษาไทย
    labels_thai = ["น้ำ", "ปูนซีเมนต์", "มวลรวมละเอียด", "มวลรวมหยาบ"]
    values = [
        mix["Water"],
        mix["Cement"],
        mix["Fine Aggregate"],
        mix["Coarse Aggregate"]
    ]
    
    fig, ax = plt.subplots(figsize=(8, 6))
    wedges, texts, autotexts = ax.pie(
        values,
        labels=labels_thai,
        autopct="%1.1f%%",
        startangle=90
    )
    
    # ตั้งค่าฟอนต์ Loma สำหรับภาษาไทย
    for text in texts:
        text.set_fontsize(14)
        text.set_family('Loma')
    for autotext in autotexts:
        autotext.set_color('white')
        autotext.set_fontsize(11)
        autotext.set_weight('bold')
    
    ax.axis("equal")
    st.pyplot(fig)

    # ---- Create Word Report Button ----
    st.subheader("📄 สร้างรายงาน Word แบบละเอียด")
    
    if st.button("📝 สร้างรายงาน Word"):
        with st.spinner("กำลังสร้างรายงาน..."):
            try:
                # เตรียมข้อมูลสำหรับรายงาน
                input_dict = {
                    'wc_ratio': wc_ratio,
                    'max_agg_mm': max_agg_mm,
                    'sg_cement': sg_cement,
                    'sg_fine': sg_fine,
                    'sg_coarse': sg_coarse,
                    'air_content': air_content,
                    'unit_weight_coarse': unit_weight_coarse,
                    'mc_fine': mc_fine,
                    'abs_fine': abs_fine,
                    'mc_coarse': mc_coarse,
                    'abs_coarse': abs_coarse
                }
                
                moisture_dict = {
                    'dw_fine': dw_fine,
                    'batch_fine': batch_fine,
                    'dw_coarse': dw_coarse,
                    'batch_coarse': batch_coarse,
                    'total_delta_water': total_delta_water,
                    'corrected_water': corrected_water
                }
                
                # สร้างรายงาน
                report_path = create_word_report(input_dict, mix, moisture_dict)
                
                # อ่านไฟล์เพื่อดาวน์โหลด
                with open(report_path, 'rb') as f:
                    report_bytes = f.read()
                
                st.success("✅ สร้างรายงานสำเร็จ!")
                
                st.download_button(
                    label="📥 ดาวน์โหลดรายงาน Word",
                    data=report_bytes,
                    file_name="รายงานออกแบบส่วนผสมคอนกรีต_ACI211.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:
                st.error(f"❌ เกิดข้อผิดพลาด: {e}")
                import traceback
                st.code(traceback.format_exc())

    st.success("คำนวณเรียบร้อย ✔")

# ---------------- Footer ----------------
st.markdown("---")
st.caption(
    "🎓 เทคโนโลยีคอนกรีต | การออกแบบส่วนผสมตามมาตรฐาน ACI 211.1 | "
    "รองรับการสอน ป.ตรี–โท และงานหน้างานเบื้องต้น"
)
